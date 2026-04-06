"""
Generador de documentos para Boletín Eléctrico de Vehículo.
Genera:
  - MTD (Memoria Técnica de Diseño) — .docx rellenado desde plantilla Word
  - CIE (Certificado de Instalación Eléctrica) — rellenado sobre el PDF oficial AcroForm
"""
import argparse
import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

import pypdf

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent

MTD_TEMPLATE_PATH = Path(
    r"\\PhicanServer\Estudio\2025\06 SERVICIOS"
    r"\S-064-2025-3584 KTP\5. DOCUMENTACIÓN\MTD_064_2025.docx"
)
FALLBACK_TEMPLATE = BASE_DIR / "PLANTILLA_MTD_VEHICULO.docx"

CIE_TEMPLATE_PATH = Path(
    r"\\PhicanServer\Estudio\2025\06 SERVICIOS"
    r"\S-064-2025-3584 KTP\5. DOCUMENTACIÓN\CIE_64-2025.pdf"
)
FALLBACK_CIE_TEMPLATE = BASE_DIR / "PLANTILLA_CIE_VEHICULO.pdf"

CONFIG_PATH = BASE_DIR / "config.json"

# ── Table cell map: (tabla_idx, row_idx, col_idx) → datos key ─────────────────
TABLE_MAP = {
    # Titular (tabla 0)
    (0, 0, 1): "TITULAR_NOMBRE",
    (0, 1, 1): "TITULAR_APELLIDOS",
    (0, 2, 1): "TITULAR_DNI",
    (0, 3, 1): "TITULAR_TELEFONO",
    (0, 4, 1): "TITULAR_EMAIL",
    # Vehículo (tabla 1)
    (1, 0, 1): "VEH_MARCA",
    (1, 1, 1): "VEH_TIPO",
    (1, 2, 1): "VEH_DENOMINACION",
    (1, 3, 1): "VEH_BASTIDOR",
    (1, 4, 1): "VEH_MATRICULA",
    # CIE asociado (tabla 2)
    (2, 0, 1): "CIE_REF",
    # Características generales (tabla 3)
    (3, 0, 1): "POT_NOMINAL",
    (3, 1, 1): "POT_PICO_CAMPO",
    (3, 2, 1): "TENSION_NOMINAL",
    # Generador FV (tabla 4)
    (4, 0, 1): "GEN_FABRICANTE",
    (4, 1, 1): "GEN_MODELO",
    (4, 2, 1): "GEN_POT_PICO",
    (4, 3, 1): "GEN_NUM_MODULOS",
    # Conductores CC (tabla 7, col 1)
    (7, 1, 1): "COND_NATURALEZA_CC",
    (7, 2, 1): "COND_AISLAMIENTO_CC",
    (7, 3, 1): "COND_CLASE_CC",
    (7, 4, 1): "COND_SECCION_CC",
    # Receptores (tabla 8)
    (8, 2, 0): "ALUM_DENOMINACION",
    (8, 2, 1): "ALUM_POTENCIA",
    (8, 2, 3): "FUERZA_DENOMINACION",
    (8, 2, 4): "FUERZA_POTENCIA",
    (8, 5, 4): "POT_TOTAL",
    # Cálculos (tabla 9)
    (9, 2, 1): "CALC_LONGITUD",
    (9, 3, 1): "CALC_MATERIAL",
    (9, 4, 1): "CALC_INTENSIDAD",
    (9, 5, 1): "CALC_CAIDA_PCT",
    (9, 5, 3): "CALC_CAIDA_V",
    (9, 6, 1): "CALC_TENSION",
    (9, 8, 1): "CALC_SECCION",
    # Técnico (tabla 10)
    (10, 0, 1): "TEC_NOMBRE",
    (10, 1, 1): "TEC_APELLIDOS",
    (10, 2, 1): "TEC_DOMICILIO",
    (10, 3, 1): "TEC_TELEFONO",
    (10, 4, 1): "TEC_EMAIL",
    (10, 5, 1): "TEC_NUM_COLEGIADO",
    (10, 6, 1): "TEC_COLEGIO",
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_text(cell, value: str):
    """Escribe value en la celda preservando el formato del primer run existente."""
    if value is None:
        value = ""
    value = str(value)

    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    # Intentar preservar formato del primer run
    if para.runs:
        run = para.runs[0]
        # Guardar propiedades del run original
        bold = run.bold
        italic = run.italic
        font_size = run.font.size
        font_name = run.font.name

        # Limpiar todos los runs
        for r in para.runs:
            r.text = ""
        para.runs[0].text = value

        # Restaurar formato
        run = para.runs[0]
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if font_size:
            run.font.size = font_size
        if font_name:
            run.font.name = font_name
    else:
        # Celda vacía — simplemente añadir run con el valor
        run = para.add_run(value)

    # Limpiar runs adicionales vacíos (restos)
    for r in list(para.runs)[1:]:
        r.text = ""


def _border_table(table):
    """Aplica bordes simples a todas las celdas de la tabla."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def _shade_cell(cell, fill_hex: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _heading_cell(cell, text: str, font_size: int = 9):
    """Escribe texto en negrita y con fondo gris en una celda de encabezado."""
    _shade_cell(cell, "D9D9D9")
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)


def _label_value_row(table, label: str, value: str, label_col: int = 0, value_col: int = 1,
                     row_idx: int = None):
    """Escribe label y value en la fila indicada de la tabla (usa add_row si row_idx=None)."""
    if row_idx is not None:
        row = table.rows[row_idx]
    else:
        row = table.add_row()
    row.cells[label_col].text = label
    row.cells[label_col].paragraphs[0].runs[0].bold = True if row.cells[label_col].paragraphs[0].runs else False
    _set_cell_text(row.cells[value_col], value)
    return row


# ── MTD ───────────────────────────────────────────────────────────────────────

def generar_mtd(datos: dict, output_dir: Path) -> Path:
    """
    Rellena la plantilla MTD con los datos proporcionados.
    Intenta abrir la plantilla en red; si falla usa la local (o crea desde cero).
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = None
    # Intentar plantilla en red
    if MTD_TEMPLATE_PATH.exists():
        try:
            doc = Document(str(MTD_TEMPLATE_PATH))
        except Exception as e:
            print(f"[MTD] No se pudo abrir plantilla de red: {e}", file=sys.stderr)

    # Fallback: plantilla local
    if doc is None and FALLBACK_TEMPLATE.exists():
        try:
            doc = Document(str(FALLBACK_TEMPLATE))
        except Exception as e:
            print(f"[MTD] No se pudo abrir plantilla local: {e}", file=sys.stderr)

    # Último recurso: documento vacío con estructura mínima
    if doc is None:
        print("[MTD] Usando documento vacío como base.", file=sys.stderr)
        doc = _crear_mtd_base(datos)
        output_name = f"MTD_{datos.get('CIE_NUMERO', 'BOLETIN')}.docx"
        out = output_dir / output_name
        doc.save(str(out))
        return out

    # Asegurar CIE_REF si no está en datos pero sí CIE_NUMERO
    if "CIE_REF" not in datos and datos.get("CIE_NUMERO"):
        datos = dict(datos)
        datos["CIE_REF"] = f"CIE: {datos['CIE_NUMERO']}"

    tables = doc.tables

    for (t_idx, r_idx, c_idx), key in TABLE_MAP.items():
        value = datos.get(key, "")
        if value is None:
            value = ""
        value = str(value)

        try:
            table = tables[t_idx]
        except IndexError:
            continue

        try:
            row = table.rows[r_idx]
        except IndexError:
            continue

        try:
            cell = row.cells[c_idx]
        except IndexError:
            continue

        _set_cell_text(cell, value)

    # Párrafo de cierre con la fecha
    lugar = datos.get("LUGAR_FIRMA", "Santa Úrsula")
    dia = datos.get("DIA_FIRMA", "")
    mes = datos.get("MES_FIRMA", "")
    anio = datos.get("ANIO_FIRMA", "")
    firma_text = f"En {lugar} a {dia} de {mes} de {anio}"

    for para in doc.paragraphs:
        text_lower = para.text.lower()
        if "en santa" in text_lower or ("en " in text_lower and " a " in text_lower and " de " in text_lower):
            # Preservar formato del primer run
            if para.runs:
                bold = para.runs[0].bold
                italic = para.runs[0].italic
                fsize = para.runs[0].font.size
                fname = para.runs[0].font.name
                for r in para.runs:
                    r.text = ""
                para.runs[0].text = firma_text
                if bold is not None:
                    para.runs[0].bold = bold
                if fsize:
                    para.runs[0].font.size = fsize
                if fname:
                    para.runs[0].font.name = fname
            else:
                para.clear()
                para.add_run(firma_text)
            break

    output_name = f"MTD_{datos.get('CIE_NUMERO', 'BOLETIN')}.docx"
    out = output_dir / output_name
    doc.save(str(out))
    return out


def _crear_mtd_base(datos: dict) -> Document:
    """Crea un documento MTD desde cero con la estructura mínima esperada."""
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    doc.add_heading("MEMORIA TÉCNICA DE DISEÑO", level=1)
    doc.add_heading("Instalación Eléctrica en Vehículo", level=2)

    def _seccion(titulo, campos):
        doc.add_heading(titulo, level=3)
        t = doc.add_table(rows=len(campos), cols=2)
        t.style = "Table Grid"
        for i, (lbl, key) in enumerate(campos):
            t.rows[i].cells[0].text = lbl
            _set_cell_text(t.rows[i].cells[1], str(datos.get(key, "")))

    _seccion("Titular", [
        ("NOMBRE", "TITULAR_NOMBRE"),
        ("APELLIDOS", "TITULAR_APELLIDOS"),
        ("DNI", "TITULAR_DNI"),
        ("TELÉFONO", "TITULAR_TELEFONO"),
        ("CORREO-E", "TITULAR_EMAIL"),
    ])
    _seccion("Vehículo", [
        ("MARCA", "VEH_MARCA"),
        ("TIPO", "VEH_TIPO"),
        ("DENOMINACIÓN COMERCIAL", "VEH_DENOMINACION"),
        ("NÚMERO DE BASTIDOR", "VEH_BASTIDOR"),
        ("MATRÍCULA", "VEH_MATRICULA"),
    ])

    doc.add_paragraph()
    lugar = datos.get("LUGAR_FIRMA", "Santa Úrsula")
    dia = datos.get("DIA_FIRMA", "")
    mes = datos.get("MES_FIRMA", "")
    anio = datos.get("ANIO_FIRMA", "")
    doc.add_paragraph(f"En {lugar} a {dia} de {mes} de {anio}")

    return doc


# ── CIE ───────────────────────────────────────────────────────────────────────

def generar_cie(datos: dict, output_dir: Path) -> Path:
    """
    Rellena el formulario oficial CIE (PDF con AcroForm) con los datos proporcionados.
    IMPORTANTE: Se cumplimenta sobre el PDF oficial, no se crea un documento nuevo.
    El PDF original tiene 84 campos AcroForm (Text1-Text84, Button19-Button69).
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # ── Resolver plantilla PDF ─────────────────────────────────────────────────
    template_path = None
    if CIE_TEMPLATE_PATH.exists():
        template_path = CIE_TEMPLATE_PATH
        # Guardar copia local como fallback offline
        if not FALLBACK_CIE_TEMPLATE.exists():
            try:
                shutil.copy2(str(CIE_TEMPLATE_PATH), str(FALLBACK_CIE_TEMPLATE))
            except Exception as e:
                print(f"[CIE] No se pudo copiar plantilla local: {e}", file=sys.stderr)
    elif FALLBACK_CIE_TEMPLATE.exists():
        template_path = FALLBACK_CIE_TEMPLATE

    if template_path is None:
        raise FileNotFoundError(
            "No se encontró la plantilla PDF del CIE. Conecta al servidor PhicanServer "
            f"o copia el archivo CIE_64-2025.pdf como: {FALLBACK_CIE_TEMPLATE}"
        )

    # ── Construir el diccionario de campos AcroForm ────────────────────────────
    titular_completo = (
        datos.get("TITULAR_NOMBRE", "") + " " + datos.get("TITULAR_APELLIDOS", "")
    ).strip()
    instalador_completo = (
        datos.get("TEC_NOMBRE", "") + " " + datos.get("TEC_APELLIDOS", "")
    ).strip()

    # Observaciones: si hay claves _L1/_L2/_L3/_L4 se usan; si no, la general en L1
    obs_general = datos.get("OBSERVACIONES", "")
    obs_l1 = datos.get("OBSERVACIONES_L1") or obs_general
    obs_l2 = datos.get("OBSERVACIONES_L2", "")
    obs_l3 = datos.get("OBSERVACIONES_L3", "")
    obs_l4 = datos.get("OBSERVACIONES_L4", "")

    text_fields = {
        # Titular / peticionario
        "Text1":  titular_completo,
        "Text2":  datos.get("TITULAR_DNI", ""),
        # Emplazamiento
        "Text3":  datos.get("EMPLAZAMIENTO_DIRECCION", ""),
        "Text38": datos.get("EMPLAZAMIENTO_NUM", ""),
        "Text5":  datos.get("EMPLAZAMIENTO_PORTAL", ""),
        "Text11": datos.get("EMPLAZAMIENTO_TM", ""),
        "Text14": datos.get("TITULAR_TELEFONO", ""),
        "Text39": datos.get("EMPLAZAMIENTO_CP", ""),
        "Text12": datos.get("EMPLAZAMIENTO_ISLA", "Tenerife"),
        "Text4":  datos.get("EMPLAZAMIENTO_USO", "Vehículo"),
        "Text40": datos.get("EMPLAZAMIENTO_SUPERFICIE", ""),
        "Text6":  datos.get("EMPLAZAMIENTO_PLANTAS", ""),
        # Potencias
        "Text10": datos.get("P_PREVISTA", ""),
        "Text16": datos.get("P_INSTALADA", ""),
        "Text7":  datos.get("P_CONTRATADA", ""),
        "Text13": datos.get("TENSION", ""),
        # Protecciones — valores de corriente / tensión
        "Text70": datos.get("PROT_IGA", ""),
        "Text72": datos.get("PROT_MAGNETO", ""),
        "Text71": datos.get("PROT_SOBRETENCION", ""),
        "Text73": datos.get("PROT_DIFERENCIAL", ""),
        # Protecciones — Icc / sensibilidad / categoría
        "Text74": datos.get("PROT_IGA_ICC", ""),
        "Text77": datos.get("PROT_MAGNETO_ICC", ""),
        "Text75": datos.get("PROT_SOBRETENCION_CAT", ""),
        "Text76": datos.get("PROT_DIFERENCIAL_MA", ""),
        # Conductores
        "Text55": datos.get("DERIVACION_CU", ""),
        "Text56": datos.get("ACOMETIDA_BT", ""),
        "Text57": datos.get("LINEA_GENERAL", ""),
        # Mediciones
        "Text17": datos.get("MED_PAT", ""),
        "Text18": datos.get("MED_AISLAMIENTO", ""),
        # Observaciones (4 líneas en el PDF)
        "Text54": obs_l1,
        "Text82": obs_l2,
        "Text83": obs_l3,
        "Text84": obs_l4,
        # Instalador / técnico director
        "Text45": instalador_completo,
        "Text48": datos.get("TEC_NUM_COLEGIADO", ""),
        "Text46": datos.get("TEC_COLEGIO", ""),
        "Text47": datos.get("TEC_NUM_COLEGIADO", ""),
        "Text49": datos.get("TEC_EMAIL", ""),
        "Text50": datos.get("TEC_TELEFONO", ""),
        # Empresa distribuidora / comercializadora
        "Text51": datos.get("EMP_DISTRIBUIDORA", ""),
        "Text52": "",
        "Text53": "",
        # Nº instalación (campo libre del formulario)
        "Text19": datos.get("CIE_NUMERO", ""),
        # Fecha y lugar de firma
        "Text78": datos.get("LUGAR_FIRMA", ""),
        "Text80": datos.get("DIA_FIRMA", ""),
        "Text81": datos.get("MES_FIRMA", ""),
        "Text79": datos.get("ANIO_FIRMA", ""),
    }

    # ── Checkboxes / Radio-buttons ─────────────────────────────────────────────
    # Objetivo de la instalación (Button26=nueva, 27=modificación, 28=ampliación, 29=cambio tensión)
    objetivo = datos.get("OBJETIVO", "Instalación nueva")
    btn_objetivo = {
        "Button26": "/No",   "Button27": "/Off",
        "Button28": "/Off",  "Button29": "/Off",
    }
    if "odificaci" in objetivo:      # Modificación
        btn_objetivo = {"Button26": "/Off", "Button27": "/No",  "Button28": "/Off", "Button29": "/Off"}
    elif "mpliaci" in objetivo:      # Ampliación
        btn_objetivo = {"Button26": "/Off", "Button27": "/Off", "Button28": "/No",  "Button29": "/Off"}
    elif "ambio" in objetivo:        # Cambio de tensión
        btn_objetivo = {"Button26": "/Off", "Button27": "/Off", "Button28": "/Off", "Button29": "/No"}

    # Documentos técnicos (Button33=Proyecto, 34=MTD, 35=Esquema, 36=Memoria, 37=Otro)
    docs = datos.get("DOCS_TECNICOS", "MTD")
    btn_docs = {f"Button{n}": "/Off" for n in range(33, 38)}
    if "MTD" in str(docs):
        btn_docs["Button34"] = "/No"
    elif "Proyecto" in str(docs):
        btn_docs["Button33"] = "/No"

    # Categoría del instalador (Button19=Básica, 20=E1, 21=E6, 22=E7, 23=E8, 24=E9, 25=otro)
    btn_cat = {f"Button{n}": "/Off" for n in range(19, 26)}
    btn_cat["Button19"] = "/No"   # Básica por defecto

    # Verificaciones UNE 20 460 (Button60-69) — dejar en /Off por defecto
    btn_une = {f"Button{n}": "/Off" for n in range(60, 70)}

    all_fields = {**text_fields, **btn_objetivo, **btn_docs, **btn_cat, **btn_une}

    # ── Rellenar PDF y guardar ────────────────────────────────────────────────
    reader = pypdf.PdfReader(str(template_path))
    writer = pypdf.PdfWriter()
    writer.append(reader)
    writer.update_page_form_field_values(writer.pages[0], all_fields)

    output_name = f"CIE_{datos.get('CIE_NUMERO', 'BOLETIN')}.pdf"
    out = output_dir / output_name
    with open(str(out), "wb") as fh:
        writer.write(fh)

    return out


def _cie_legacy_docx(datos: dict, output_dir: Path) -> Path:
    """
    Genera una versión DOCX del CIE como documento de respaldo (sin valor oficial).
    Solo se usa si el PDF oficial no está disponible y se necesita un borrador.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Configuración de página ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Estilo base: tamaño de fuente pequeño para caber en una página
    style = doc.styles["Normal"]
    style.font.size = Pt(8)
    style.font.name = "Arial"

    def _para(text="", bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Arial"
        return p

    def _cell_label(cell, text, size=7, bold=True):
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"

    def _cell_value(cell, text, size=8, bold=False):
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if para.runs:
            para.runs[0].text = str(text) if text else ""
        else:
            run = para.add_run(str(text) if text else "")
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Arial"

    # ──────────────────────────────────────────────────────────────────────────
    # ENCABEZADO
    # ──────────────────────────────────────────────────────────────────────────
    t_header = doc.add_table(rows=2, cols=3)
    t_header.style = "Table Grid"
    t_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Col widths: logo | título | nº instalación + expediente
    t_header.columns[0].width = Cm(4)
    t_header.columns[1].width = Cm(9)
    t_header.columns[2].width = Cm(5)

    # Celda 0: Logo/empresa
    _shade_cell(t_header.rows[0].cells[0], "1F497D")
    p_logo = t_header.rows[0].cells[0].paragraphs[0]
    p_logo.paragraph_format.space_before = Pt(4)
    p_logo.paragraph_format.space_after = Pt(4)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PHICAN INGENIEROS")
    r_logo.bold = True
    r_logo.font.size = Pt(10)
    r_logo.font.name = "Arial"
    r_logo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Celda 1: Título
    _shade_cell(t_header.rows[0].cells[1], "1F497D")
    p_tit = t_header.rows[0].cells[1].paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_before = Pt(4)
    p_tit.paragraph_format.space_after = Pt(2)
    r_tit = p_tit.add_run("CERTIFICADO DE INSTALACIÓN ELÉCTRICA")
    r_tit.bold = True
    r_tit.font.size = Pt(11)
    r_tit.font.name = "Arial"
    r_tit.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_tit2 = t_header.rows[0].cells[1].add_paragraph("C.I. B.T.")
    p_tit2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit2.runs[0].bold = True
    p_tit2.runs[0].font.size = Pt(9)
    p_tit2.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Celda 2: Nº instalación / expediente
    c_nums = t_header.rows[0].cells[2]
    p_n1 = c_nums.paragraphs[0]
    r_n1 = p_n1.add_run("Nº de Instalación: ")
    r_n1.bold = True
    r_n1.font.size = Pt(8)
    p_n1.add_run(str(datos.get("CIE_NUMERO", "")))
    p_n2 = c_nums.add_paragraph()
    r_n2 = p_n2.add_run("Nº de Expediente: ")
    r_n2.bold = True
    r_n2.font.size = Pt(8)
    p_n2.add_run(str(datos.get("CIE_EXPEDIENTE", "")))

    # Fila 2: BAJA TENSIÓN banner
    baja = t_header.rows[1].cells[0]
    # Merge las 3 celdas de la segunda fila
    baja.merge(t_header.rows[1].cells[2])
    _shade_cell(baja, "4472C4")
    p_baja = baja.paragraphs[0]
    p_baja.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_baja = p_baja.add_run("BAJA TENSIÓN")
    r_baja.bold = True
    r_baja.font.size = Pt(10)
    r_baja.font.name = "Arial"
    r_baja.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ──────────────────────────────────────────────────────────────────────────
    # TITULAR
    # ──────────────────────────────────────────────────────────────────────────
    _para()  # espacio
    t_titular = doc.add_table(rows=3, cols=4)
    t_titular.style = "Table Grid"
    _border_table(t_titular)

    # Row 0: sección header
    t_titular.rows[0].cells[0].merge(t_titular.rows[0].cells[3])
    _shade_cell(t_titular.rows[0].cells[0], "4472C4")
    _cell_label(t_titular.rows[0].cells[0], "TITULAR / PETICIONARIO", size=9, bold=True)
    tc_pct = t_titular.rows[0].cells[0]._tc
    _WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for para in tc_pct.iter(f"{{{_WNS}}}p"):
        for run_el in para.iter(f"{{{_WNS}}}r"):
            rPr = run_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run_el.insert(0, rPr)
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), "FFFFFF")
            rPr.append(color_el)

    # Row 1: Nombre, DOI
    _cell_label(t_titular.rows[1].cells[0], "Nombre / Razón social:")
    _cell_value(t_titular.rows[1].cells[1], datos.get("TITULAR_NOMBRE", "") + " " + datos.get("TITULAR_APELLIDOS", ""))
    _cell_label(t_titular.rows[1].cells[2], "D.O.I. / C.I.F.:")
    _cell_value(t_titular.rows[1].cells[3], datos.get("TITULAR_DNI", ""))

    # Row 2: Teléfono, Email
    _cell_label(t_titular.rows[2].cells[0], "Teléfono:")
    _cell_value(t_titular.rows[2].cells[1], datos.get("TITULAR_TELEFONO", ""))
    _cell_label(t_titular.rows[2].cells[2], "Correo-e:")
    _cell_value(t_titular.rows[2].cells[3], datos.get("TITULAR_EMAIL", ""))

    # ──────────────────────────────────────────────────────────────────────────
    # EMPLAZAMIENTO
    # ──────────────────────────────────────────────────────────────────────────
    t_emp = doc.add_table(rows=4, cols=6)
    t_emp.style = "Table Grid"
    _border_table(t_emp)

    # Row 0: header
    t_emp.rows[0].cells[0].merge(t_emp.rows[0].cells[5])
    _shade_cell(t_emp.rows[0].cells[0], "4472C4")
    _cell_label(t_emp.rows[0].cells[0], "EMPLAZAMIENTO / VEHÍCULO", size=9)
    _set_header_white(t_emp.rows[0].cells[0])

    # Row 1: Dirección | N° | Portal | T.M.
    _cell_label(t_emp.rows[1].cells[0], "Dirección / Vehículo:")
    t_emp.rows[1].cells[1].merge(t_emp.rows[1].cells[2])
    _cell_value(t_emp.rows[1].cells[1], datos.get("EMPLAZAMIENTO_DIRECCION", datos.get("VEH_MARCA", "") + " " + datos.get("VEH_MATRICULA", "")))
    _cell_label(t_emp.rows[1].cells[3], "Matrícula:")
    _cell_value(t_emp.rows[1].cells[4], datos.get("VEH_MATRICULA", ""))
    _cell_label(t_emp.rows[1].cells[5], "Bastidor:")

    # Row 2: Bastidor / municipio
    _cell_label(t_emp.rows[2].cells[0], "Nº Bastidor:")
    t_emp.rows[2].cells[1].merge(t_emp.rows[2].cells[2])
    _cell_value(t_emp.rows[2].cells[1], datos.get("VEH_BASTIDOR", ""))
    _cell_label(t_emp.rows[2].cells[3], "T.M.:")
    _cell_value(t_emp.rows[2].cells[4], datos.get("EMPLAZAMIENTO_TM", ""))
    _cell_label(t_emp.rows[2].cells[5], "Isla:")

    # Row 3: Teléfono, CP
    _cell_label(t_emp.rows[3].cells[0], "Tfno.:")
    _cell_value(t_emp.rows[3].cells[1], datos.get("TITULAR_TELEFONO", ""))
    _cell_label(t_emp.rows[3].cells[2], "C.P.:")
    _cell_value(t_emp.rows[3].cells[3], datos.get("EMPLAZAMIENTO_CP", ""))
    _cell_label(t_emp.rows[3].cells[4], "Uso / Destino:")
    _cell_value(t_emp.rows[3].cells[5], datos.get("EMPLAZAMIENTO_USO", "Vehículo"))

    # ──────────────────────────────────────────────────────────────────────────
    # PRINCIPALES CARACTERÍSTICAS ELÉCTRICAS
    # ──────────────────────────────────────────────────────────────────────────
    t_caract = doc.add_table(rows=1, cols=2)
    t_caract.style = "Table Grid"
    _border_table(t_caract)

    # Sub-tabla izquierda: POTENCIAS
    left_cell = t_caract.rows[0].cells[0]
    right_cell = t_caract.rows[0].cells[1]

    # Header row (fusioned)
    header_row = t_caract.rows[0]
    header_row.cells[0].merge(header_row.cells[1])
    _shade_cell(header_row.cells[0], "4472C4")
    _cell_label(header_row.cells[0], "PRINCIPALES CARACTERÍSTICAS ELÉCTRICAS", size=9)
    _set_header_white(header_row.cells[0])

    # Nueva tabla anidada para potencias / protecciones
    t_potprot = doc.add_table(rows=6, cols=4)
    t_potprot.style = "Table Grid"
    _border_table(t_potprot)

    # Sub-headers
    t_potprot.rows[0].cells[0].merge(t_potprot.rows[0].cells[1])
    _shade_cell(t_potprot.rows[0].cells[0], "D9D9D9")
    _cell_label(t_potprot.rows[0].cells[0], "POTENCIAS", size=8)
    t_potprot.rows[0].cells[2].merge(t_potprot.rows[0].cells[3])
    _shade_cell(t_potprot.rows[0].cells[2], "D9D9D9")
    _cell_label(t_potprot.rows[0].cells[2], "PROTECCIONES", size=8)

    potencias = [
        ("P. prevista", datos.get("P_PREVISTA", ""), "I.G.A.", datos.get("PROT_IGA", "")),
        ("P. instalada", datos.get("P_INSTALADA", ""), "Magnetotérmicos", datos.get("PROT_MAGNETO", "")),
        ("P. contratada recomendada", datos.get("P_CONTRATADA", ""), "Sobretensiones/cat.", datos.get("PROT_SOBRETENCION", "")),
        ("Tensión", datos.get("TENSION", datos.get("TENSION_NOMINAL", "")), "Diferencial/sensib.", datos.get("PROT_DIFERENCIAL", "")),
        ("Der. Individual Cu/Al", datos.get("DERIVACION_CU", ""), "Resist. p.a.t. Ω", datos.get("MED_PAT", "")),
        ("Acometida Red BT", datos.get("ACOMETIDA_BT", ""), "Resist. aislamiento KΩ", datos.get("MED_AISLAMIENTO", "")),
    ]

    for i, (lbl_p, val_p, lbl_pr, val_pr) in enumerate(potencias):
        row = t_potprot.rows[i + 1] if i + 1 < len(t_potprot.rows) else t_potprot.add_row()
        _cell_label(row.cells[0], lbl_p + ":", size=7)
        _cell_value(row.cells[1], val_p, size=8)
        _cell_label(row.cells[2], lbl_pr + ":", size=7)
        _cell_value(row.cells[3], val_pr, size=8)

    # ──────────────────────────────────────────────────────────────────────────
    # OBJETIVO + DOCUMENTOS TÉCNICOS
    # ──────────────────────────────────────────────────────────────────────────
    t_obj = doc.add_table(rows=3, cols=4)
    t_obj.style = "Table Grid"
    _border_table(t_obj)

    # Header
    t_obj.rows[0].cells[0].merge(t_obj.rows[0].cells[3])
    _shade_cell(t_obj.rows[0].cells[0], "4472C4")
    _cell_label(t_obj.rows[0].cells[0], "OBJETIVO / DOCUMENTOS TÉCNICOS", size=9)
    _set_header_white(t_obj.rows[0].cells[0])

    # Objetivo checkboxes (texto aproximado)
    objetivo = datos.get("OBJETIVO", "Instalación nueva")
    _cell_label(t_obj.rows[1].cells[0], "Objetivo:")
    _cell_value(t_obj.rows[1].cells[1], objetivo)
    _cell_label(t_obj.rows[1].cells[2], "Documentos técnicos:")
    _cell_value(t_obj.rows[1].cells[3], datos.get("DOCS_TECNICOS", "MTD"))

    _cell_label(t_obj.rows[2].cells[0], "Verificaciones UNE 20 460-6-61:")
    _cell_value(t_obj.rows[2].cells[1], "Sí")
    _cell_label(t_obj.rows[2].cells[2], "Empresa distribuidora:")
    _cell_value(t_obj.rows[2].cells[3], datos.get("EMP_DISTRIBUIDORA", ""))

    # ──────────────────────────────────────────────────────────────────────────
    # INSTALADOR
    # ──────────────────────────────────────────────────────────────────────────
    t_inst = doc.add_table(rows=4, cols=4)
    t_inst.style = "Table Grid"
    _border_table(t_inst)

    t_inst.rows[0].cells[0].merge(t_inst.rows[0].cells[3])
    _shade_cell(t_inst.rows[0].cells[0], "4472C4")
    _cell_label(t_inst.rows[0].cells[0], "INSTALADOR / TÉCNICO DIRECTOR", size=9)
    _set_header_white(t_inst.rows[0].cells[0])

    _cell_label(t_inst.rows[1].cells[0], "Nombre:")
    _cell_value(t_inst.rows[1].cells[1],
                datos.get("TEC_NOMBRE", "") + " " + datos.get("TEC_APELLIDOS", ""))
    _cell_label(t_inst.rows[1].cells[2], "Nº C.C.I. / Colegiado:")
    _cell_value(t_inst.rows[1].cells[3], datos.get("TEC_NUM_COLEGIADO", ""))

    _cell_label(t_inst.rows[2].cells[0], "Empresa / Colegio:")
    _cell_value(t_inst.rows[2].cells[1], datos.get("TEC_COLEGIO", ""))
    _cell_label(t_inst.rows[2].cells[2], "Tfno.:")
    _cell_value(t_inst.rows[2].cells[3], datos.get("TEC_TELEFONO", ""))

    _cell_label(t_inst.rows[3].cells[0], "Email:")
    _cell_value(t_inst.rows[3].cells[1], datos.get("TEC_EMAIL", ""))
    _cell_label(t_inst.rows[3].cells[2], "Domicilio:")
    _cell_value(t_inst.rows[3].cells[3], datos.get("TEC_DOMICILIO", ""))

    # ──────────────────────────────────────────────────────────────────────────
    # FECHA Y FIRMA
    # ──────────────────────────────────────────────────────────────────────────
    t_fecha = doc.add_table(rows=2, cols=3)
    t_fecha.style = "Table Grid"
    _border_table(t_fecha)

    t_fecha.rows[0].cells[0].merge(t_fecha.rows[0].cells[2])
    _shade_cell(t_fecha.rows[0].cells[0], "4472C4")
    _cell_label(t_fecha.rows[0].cells[0], "FECHA Y FIRMA", size=9)
    _set_header_white(t_fecha.rows[0].cells[0])

    lugar = datos.get("LUGAR_FIRMA", "Santa Úrsula")
    dia = datos.get("DIA_FIRMA", "")
    mes = datos.get("MES_FIRMA", "")
    anio = datos.get("ANIO_FIRMA", "")
    _cell_label(t_fecha.rows[1].cells[0], "Lugar:")
    _cell_value(t_fecha.rows[1].cells[0], lugar)
    _cell_label(t_fecha.rows[1].cells[1], "Fecha:")
    _cell_value(t_fecha.rows[1].cells[1], f"{dia} de {mes} de {anio}")
    _cell_label(t_fecha.rows[1].cells[2], "Firma:")
    _cell_value(t_fecha.rows[1].cells[2], "")

    # ──────────────────────────────────────────────────────────────────────────
    # OBSERVACIONES
    # ──────────────────────────────────────────────────────────────────────────
    t_obs = doc.add_table(rows=2, cols=1)
    t_obs.style = "Table Grid"
    _border_table(t_obs)

    t_obs.rows[0].cells[0]
    _shade_cell(t_obs.rows[0].cells[0], "4472C4")
    _cell_label(t_obs.rows[0].cells[0], "OBSERVACIONES", size=9)
    _set_header_white(t_obs.rows[0].cells[0])

    obs_cell = t_obs.rows[1].cells[0]
    obs_text = datos.get("OBSERVACIONES", "")
    obs_para = obs_cell.paragraphs[0]
    obs_para.paragraph_format.space_before = Pt(2)
    obs_para.paragraph_format.space_after = Pt(2)
    obs_run = obs_para.add_run(obs_text)
    obs_run.font.size = Pt(8)
    obs_run.font.name = "Arial"

    # Espacio mínimo para la observación
    obs_cell._tc.get_or_add_tcPr()

    output_name = f"CIE_BORRADOR_{datos.get('CIE_NUMERO', 'BOLETIN')}.docx"
    out = output_dir / output_name
    doc.save(str(out))
    return out


def _set_header_white(cell):
    """Pone el texto de la celda (primer párrafo) en blanco para cabeceras azules."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ── Integración con DB ────────────────────────────────────────────────────────

def _load_config() -> dict:
    """Carga config.json del proyecto base."""
    if CONFIG_PATH.exists():
        with open(CONFIG_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _default_datos(boletin_row) -> dict:
    """
    Construye el dict de datos a partir de la fila de BD y config.json.
    boletin_row puede ser un dict o sqlite3.Row.
    """
    if not isinstance(boletin_row, dict):
        boletin_row = dict(boletin_row)

    config = _load_config()

    # Parsear datos_json existente en el boletín
    datos_json = {}
    raw = boletin_row.get("datos_json") or "{}"
    if isinstance(raw, str):
        try:
            datos_json = json.loads(raw)
        except Exception:
            datos_json = {}
    elif isinstance(raw, dict):
        datos_json = raw

    # Número CIE base = número del boletín
    cie_numero = datos_json.get("CIE_NUMERO") or boletin_row.get("numero", "")

    # Titular desde campos del boletín (fallback a datos_json)
    peticionario = boletin_row.get("peticionario", "") or ""
    partes = peticionario.strip().split(" ", 1)
    titular_nombre = datos_json.get("TITULAR_NOMBRE") or (partes[0] if partes else "")
    titular_apellidos = datos_json.get("TITULAR_APELLIDOS") or (partes[1] if len(partes) > 1 else "")

    import datetime
    hoy = datetime.date.today()
    meses_es = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo",
        6: "junio", 7: "julio", 8: "agosto", 9: "septiembre",
        10: "octubre", 11: "noviembre", 12: "diciembre",
    }

    datos = {
        # CIE
        "CIE_NUMERO": cie_numero,
        "CIE_REF": f"CIE: {cie_numero}",
        "CIE_EXPEDIENTE": datos_json.get("CIE_EXPEDIENTE") or boletin_row.get("expediente_id", "") or "",
        # Titular
        "TITULAR_NOMBRE": titular_nombre,
        "TITULAR_APELLIDOS": titular_apellidos,
        "TITULAR_DNI": datos_json.get("TITULAR_DNI") or boletin_row.get("nif", "") or "",
        "TITULAR_TELEFONO": datos_json.get("TITULAR_TELEFONO") or "",
        "TITULAR_EMAIL": datos_json.get("TITULAR_EMAIL") or "",
        # Vehículo
        "VEH_MARCA": datos_json.get("VEH_MARCA", ""),
        "VEH_TIPO": datos_json.get("VEH_TIPO", ""),
        "VEH_DENOMINACION": datos_json.get("VEH_DENOMINACION", ""),
        "VEH_BASTIDOR": datos_json.get("VEH_BASTIDOR", ""),
        "VEH_MATRICULA": datos_json.get("VEH_MATRICULA", ""),
        # Instalación
        "POT_NOMINAL": datos_json.get("POT_NOMINAL", ""),
        "POT_PICO_CAMPO": datos_json.get("POT_PICO_CAMPO", ""),
        "TENSION_NOMINAL": datos_json.get("TENSION_NOMINAL", ""),
        "P_PREVISTA": datos_json.get("P_PREVISTA", ""),
        "P_INSTALADA": datos_json.get("P_INSTALADA", ""),
        "P_CONTRATADA": datos_json.get("P_CONTRATADA", ""),
        "TENSION": datos_json.get("TENSION", datos_json.get("TENSION_NOMINAL", "")),
        "DERIVACION_CU": datos_json.get("DERIVACION_CU", ""),
        "ACOMETIDA_BT": datos_json.get("ACOMETIDA_BT", ""),
        # Generador FV
        "GEN_FABRICANTE": datos_json.get("GEN_FABRICANTE", ""),
        "GEN_MODELO": datos_json.get("GEN_MODELO", ""),
        "GEN_POT_PICO": datos_json.get("GEN_POT_PICO", ""),
        "GEN_NUM_MODULOS": datos_json.get("GEN_NUM_MODULOS", ""),
        # Conductores
        "COND_NATURALEZA_CC": datos_json.get("COND_NATURALEZA_CC", ""),
        "COND_AISLAMIENTO_CC": datos_json.get("COND_AISLAMIENTO_CC", ""),
        "COND_CLASE_CC": datos_json.get("COND_CLASE_CC", ""),
        "COND_SECCION_CC": datos_json.get("COND_SECCION_CC", ""),
        # Receptores
        "ALUM_DENOMINACION": datos_json.get("ALUM_DENOMINACION", ""),
        "ALUM_POTENCIA": datos_json.get("ALUM_POTENCIA", ""),
        "FUERZA_DENOMINACION": datos_json.get("FUERZA_DENOMINACION", ""),
        "FUERZA_POTENCIA": datos_json.get("FUERZA_POTENCIA", ""),
        "POT_TOTAL": datos_json.get("POT_TOTAL", ""),
        # Cálculos
        "CALC_LONGITUD": datos_json.get("CALC_LONGITUD", ""),
        "CALC_MATERIAL": datos_json.get("CALC_MATERIAL", ""),
        "CALC_INTENSIDAD": datos_json.get("CALC_INTENSIDAD", ""),
        "CALC_CAIDA_PCT": datos_json.get("CALC_CAIDA_PCT", ""),
        "CALC_CAIDA_V": datos_json.get("CALC_CAIDA_V", ""),
        "CALC_TENSION": datos_json.get("CALC_TENSION", ""),
        "CALC_SECCION": datos_json.get("CALC_SECCION", ""),
        # Protecciones
        "PROT_IGA": datos_json.get("PROT_IGA", ""),
        "PROT_MAGNETO": datos_json.get("PROT_MAGNETO", ""),
        "PROT_SOBRETENCION": datos_json.get("PROT_SOBRETENCION", ""),
        "PROT_DIFERENCIAL": datos_json.get("PROT_DIFERENCIAL", ""),
        "MED_PAT": datos_json.get("MED_PAT", ""),
        "MED_AISLAMIENTO": datos_json.get("MED_AISLAMIENTO", ""),
        # Empresa
        "EMP_DISTRIBUIDORA": datos_json.get("EMP_DISTRIBUIDORA", ""),
        "OBJETIVO": datos_json.get("OBJETIVO", "Instalación nueva"),
        "DOCS_TECNICOS": datos_json.get("DOCS_TECNICOS", "MTD"),
        # Emplazamiento
        "EMPLAZAMIENTO_DIRECCION": datos_json.get("EMPLAZAMIENTO_DIRECCION") or boletin_row.get("direccion", "") or "",
        "EMPLAZAMIENTO_NUM": datos_json.get("EMPLAZAMIENTO_NUM", ""),
        "EMPLAZAMIENTO_PORTAL": datos_json.get("EMPLAZAMIENTO_PORTAL", ""),
        "EMPLAZAMIENTO_TM": datos_json.get("EMPLAZAMIENTO_TM") or boletin_row.get("poblacion", "") or "",
        "EMPLAZAMIENTO_CP": datos_json.get("EMPLAZAMIENTO_CP") or boletin_row.get("cp", "") or "",
        "EMPLAZAMIENTO_ISLA": datos_json.get("EMPLAZAMIENTO_ISLA") or config.get("EMPLAZAMIENTO_ISLA", "Tenerife"),
        "EMPLAZAMIENTO_USO": datos_json.get("EMPLAZAMIENTO_USO", "Vehículo"),
        "EMPLAZAMIENTO_SUPERFICIE": datos_json.get("EMPLAZAMIENTO_SUPERFICIE", ""),
        "EMPLAZAMIENTO_PLANTAS": datos_json.get("EMPLAZAMIENTO_PLANTAS", ""),
        # Protecciones — ICC y sensibilidad (campos extra del PDF CIE)
        "PROT_IGA_ICC": datos_json.get("PROT_IGA_ICC", ""),
        "PROT_MAGNETO_ICC": datos_json.get("PROT_MAGNETO_ICC", ""),
        "PROT_SOBRETENCION_CAT": datos_json.get("PROT_SOBRETENCION_CAT", ""),
        "PROT_DIFERENCIAL_MA": datos_json.get("PROT_DIFERENCIAL_MA", ""),
        # Línea general de alimentación
        "LINEA_GENERAL": datos_json.get("LINEA_GENERAL", ""),
        # Técnico (desde config.json)
        "TEC_NOMBRE": datos_json.get("TEC_NOMBRE") or config.get("TEC_NOMBRE", ""),
        "TEC_APELLIDOS": datos_json.get("TEC_APELLIDOS") or config.get("TEC_APELLIDOS", ""),
        "TEC_DOMICILIO": datos_json.get("TEC_DOMICILIO") or config.get("TEC_DOMICILIO", ""),
        "TEC_TELEFONO": datos_json.get("TEC_TELEFONO") or config.get("TEC_TELEFONO", ""),
        "TEC_EMAIL": datos_json.get("TEC_EMAIL") or config.get("TEC_EMAIL", ""),
        "TEC_NUM_COLEGIADO": datos_json.get("TEC_NUM_COLEGIADO") or config.get("TEC_NUM_COLEGIADO", ""),
        "TEC_COLEGIO": datos_json.get("TEC_COLEGIO") or config.get("TEC_COLEGIO", ""),
        # Fecha de firma
        "LUGAR_FIRMA": datos_json.get("LUGAR_FIRMA") or config.get("LUGAR_FIRMA", "Santa Úrsula"),
        "DIA_FIRMA": datos_json.get("DIA_FIRMA") or str(hoy.day),
        "MES_FIRMA": datos_json.get("MES_FIRMA") or meses_es[hoy.month],
        "ANIO_FIRMA": datos_json.get("ANIO_FIRMA") or str(hoy.year),
        # Observaciones — campo general y 4 líneas para el PDF CIE
        "OBSERVACIONES": datos_json.get("OBSERVACIONES") or boletin_row.get("observaciones", "") or "",
        "OBSERVACIONES_L1": datos_json.get("OBSERVACIONES_L1", ""),
        "OBSERVACIONES_L2": datos_json.get("OBSERVACIONES_L2", ""),
        "OBSERVACIONES_L3": datos_json.get("OBSERVACIONES_L3", ""),
        "OBSERVACIONES_L4": datos_json.get("OBSERVACIONES_L4", ""),
    }

    return datos


def generar_boletin_docs(boletin_id: int) -> dict:
    """
    Genera MTD y CIE para el boletín indicado.
    Retorna {"ok": True, "mtd": str, "cie": str} o {"ok": False, "error": str}.
    """
    try:
        # Importar db desde la carpeta tracker (si se ejecuta desde allí o desde raíz)
        try:
            from tracker.db import get_boletin  # cuando se importa desde raíz del proyecto
        except ImportError:
            import sys as _sys
            _sys.path.insert(0, str(BASE_DIR / "tracker"))
            from db import get_boletin

        bol = get_boletin(boletin_id)
        if not bol:
            return {"ok": False, "error": f"Boletín {boletin_id} no encontrado"}

        datos = _default_datos(bol)

        # Carpeta de salida: la carpeta del boletín en la BD, o una temporal
        carpeta_bol = bol.get("carpeta") or ""
        if carpeta_bol and Path(carpeta_bol).is_dir():
            output_dir = Path(carpeta_bol)
        else:
            output_dir = BASE_DIR / "BOLETINES" / "ELECTRICIDAD" / "VEHICULO" / datos["CIE_NUMERO"]
            output_dir.mkdir(parents=True, exist_ok=True)

        mtd_path = generar_mtd(datos, output_dir)
        cie_path = generar_cie(datos, output_dir)

        return {"ok": True, "mtd": str(mtd_path), "cie": str(cie_path)}

    except Exception as e:
        import traceback
        return {"ok": False, "error": str(e), "traceback": traceback.format_exc()}


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Genera documentos de Boletín Eléctrico Vehículo")
    parser.add_argument("--boletin_id", type=int, required=True,
                        help="ID del boletín en la base de datos")
    args = parser.parse_args()

    resultado = generar_boletin_docs(args.boletin_id)
    if resultado["ok"]:
        print(f"[OK] MTD: {resultado['mtd']}")
        print(f"[OK] CIE: {resultado['cie']}")
    else:
        print(f"[ERROR] {resultado['error']}", file=sys.stderr)
        if "traceback" in resultado:
            print(resultado["traceback"], file=sys.stderr)
        sys.exit(1)
