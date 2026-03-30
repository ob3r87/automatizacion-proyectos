#!/usr/bin/env python3
"""
generar_anexo.py
================
Genera un documento Word (.docx) con el Anexo Justificativo de Cálculos
a partir del JSON de datos producido por formulario.py.

Uso:
    python generar_anexo.py <ruta_json>

Salida:
    Imprime en stdout la ruta del .docx generado (para que formulario.py la capture).
"""

import sys
import json
import subprocess
from pathlib import Path
from datetime import date

# ── Auto-instalación de python-docx si no está disponible ───────────────────
try:
    from docx import Document
except ImportError:
    print("Instalando python-docx, espera un momento...", file=sys.stderr)
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "--quiet"],
        stdout=subprocess.DEVNULL
    )
    from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paleta de colores ───────────────────────────────────────────────────────
COLOR_AZUL    = RGBColor(0x19, 0x76, 0xD2)   # azul corporativo
COLOR_GRIS    = RGBColor(0x42, 0x42, 0x42)   # texto oscuro
COLOR_VERDE   = RGBColor(0x2E, 0x7D, 0x32)   # ✓ resultado OK
COLOR_ROJO    = RGBColor(0xC6, 0x28, 0x28)   # ✗ resultado KO
COLOR_AMBAR   = RGBColor(0xF5, 0x7F, 0x17)   # separadores sección
COLOR_FONDO_H = "1565C0"                      # azul oscuro encabezados tabla (hex sin #)
COLOR_FONDO_S = "E3F2FD"                      # azul muy claro filas alternas
COLOR_SEP     = "FFF8E1"                      # amarillo suave separadores intraCalculо

# ─── Constantes de layout ────────────────────────────────────────────────────
MARGEN_CM = 2.0


# ════════════════════════════════════════════════════════════════════════════
#  Utilidades de formato
# ════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda con un color hexadecimal (sin '#')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc, text, level=1, color=COLOR_AZUL):
    """Encabezado personalizado sin usar estilos de Heading del documento."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    size = {1: 14, 2: 12, 3: 10.5}.get(level, 10)
    _run(para, text, bold=True, size=size, color=color)
    return para


def _hr(doc, color_hex="BDBDBD"):
    """Línea horizontal usando borde inferior del párrafo."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.upper())
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_module.enum.text.WD_BREAK.PAGE)   # fallback below


def _add_page_break(doc):
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _col_widths(table, widths_cm):
    """Fija el ancho de cada columna en centímetros."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _tabla_dos_col(doc, filas, widths=(6.5, 10.0)):
    """Tabla de dos columnas Clave | Valor."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, valor) in enumerate(filas):
        row   = table.add_row()
        cL    = row.cells[0]
        cV    = row.cells[1]
        cL.width = Cm(widths[0])
        cV.width = Cm(widths[1])
        # Fondo alternado
        if i % 2 == 0:
            _set_cell_bg(cL, COLOR_FONDO_S)
            _set_cell_bg(cV, COLOR_FONDO_S)
        pL = cL.paragraphs[0]
        pV = cV.paragraphs[0]
        _run(pL, str(label), bold=True,  size=9, color=COLOR_GRIS)
        _run(pV, str(valor), bold=False, size=9)
    return table


# ════════════════════════════════════════════════════════════════════════════
#  Utilidades para desarrollo paso a paso de cálculos
# ════════════════════════════════════════════════════════════════════════════

COLOR_FORM = RGBColor(0x33, 0x33, 0x33)  # fórmulas

def _paso(doc, num, titulo):
    """Encabezado de paso: 'Paso N — Título'"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)
    _run(para, f"Paso {num} — ", bold=True, size=10, color=COLOR_AZUL)
    _run(para, titulo, bold=True, size=10, color=COLOR_AZUL)
    return para


def _descripcion(doc, texto):
    """Texto narrativo del paso."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Cm(0.5)
    _run(para, texto, size=9, color=COLOR_GRIS)
    return para


def _formula(doc, simbolica, sustituida=None):
    """Fórmula simbólica + opcional sustitución con valores."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Cm(1.0)
    _run(para, simbolica, italic=True, size=9.5, color=COLOR_FORM)
    if sustituida:
        para2 = doc.add_paragraph()
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after  = Pt(1)
        para2.paragraph_format.left_indent  = Cm(1.0)
        _run(para2, sustituida, size=9, color=COLOR_GRIS)
    return para


def _resultado_paso(doc, texto, ok=True):
    """Resultado final con color verde/rojo."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.5)
    color = COLOR_VERDE if ok else COLOR_ROJO
    _run(para, texto, bold=True, size=10, color=color)
    return para


def _v(campos, clave, defecto="—"):
    """Lee valor de campos, devuelve string."""
    val = campos.get(clave, "")
    return str(val).strip() if val else defecto


def _f(campos, clave, defecto=0.0):
    """Lee valor numérico de campos."""
    try:
        return float(str(campos.get(clave, "")).replace(",", ".") or defecto)
    except (ValueError, TypeError):
        return defecto


# ════════════════════════════════════════════════════════════════════════════
#  Helpers OMML — Ecuaciones matemáticas nativas de Word
# ════════════════════════════════════════════════════════════════════════════

def _omml_run(text):
    """Crea un run de texto matemático: m:r > m:t."""
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def _omml_frac(num_elems, den_elems):
    """Fracción apilada: m:f > m:num + m:den.
    num_elems y den_elems son listas de elementos OMML o strings."""
    f = OxmlElement("m:f")
    num = OxmlElement("m:num")
    den = OxmlElement("m:den")
    for el in (_omml_ensure_list(num_elems)):
        num.append(el)
    for el in (_omml_ensure_list(den_elems)):
        den.append(el)
    f.append(num)
    f.append(den)
    return f


def _omml_sub(base_text, sub_text):
    """Subíndice: m:sSub > m:e + m:sub."""
    ss = OxmlElement("m:sSub")
    e  = OxmlElement("m:e")
    e.append(_omml_run(base_text))
    ss.append(e)
    s  = OxmlElement("m:sub")
    s.append(_omml_run(sub_text))
    ss.append(s)
    return ss


def _omml_sup(base_text, sup_text):
    """Superíndice: m:sSup > m:e + m:sup."""
    ss = OxmlElement("m:sSup")
    e  = OxmlElement("m:e")
    e.append(_omml_run(base_text))
    ss.append(e)
    s  = OxmlElement("m:sup")
    s.append(_omml_run(sup_text))
    ss.append(s)
    return ss


def _omml_sqrt(content_elems):
    """Raíz cuadrada: m:rad > m:deg (vacío) + m:e."""
    rad = OxmlElement("m:rad")
    deg = OxmlElement("m:deg")
    rad.append(deg)
    e = OxmlElement("m:e")
    for el in _omml_ensure_list(content_elems):
        e.append(el)
    rad.append(e)
    return rad


def _omml_ensure_list(items):
    """Convierte strings a _omml_run, deja elementos OMML tal cual."""
    if isinstance(items, str):
        return [_omml_run(items)]
    if hasattr(items, "tag"):
        return [items]
    result = []
    for it in items:
        if isinstance(it, str):
            result.append(_omml_run(it))
        else:
            result.append(it)
    return result


def _omml_eq(paragraph, *elements):
    """Inserta una ecuación m:oMath en el párrafo con los elementos dados."""
    oMath = OxmlElement("m:oMath")
    for el in elements:
        if isinstance(el, str):
            oMath.append(_omml_run(el))
        else:
            oMath.append(el)
    paragraph._element.append(oMath)
    return paragraph


def _formula_omml(doc, omml_elements, sustituida=None):
    """Fórmula OMML (ecuación nativa de Word) + sustitución en texto plano."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Cm(1.0)
    _omml_eq(para, *omml_elements)
    if sustituida:
        para2 = doc.add_paragraph()
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after  = Pt(1)
        para2.paragraph_format.left_indent  = Cm(1.0)
        _run(para2, sustituida, size=9, color=COLOR_GRIS)
    return para


# ════════════════════════════════════════════════════════════════════════════
#  Sección de cabecera de documento
# ════════════════════════════════════════════════════════════════════════════

def _cabecera_documento(doc, datos):
    ref   = datos.get("REFERENCIA", "—")
    mat   = datos.get("MATRICULA",  "—")
    marca = datos.get("MARCA",  "")
    model = datos.get("MODELO", "")
    pet   = datos.get("PETICIONARIO_NOMBRE", "—")
    hoy   = date.today().strftime("%d/%m/%Y")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after  = Pt(6)
    _run(title, "ANEXO JUSTIFICATIVO DE CÁLCULOS",
         bold=True, size=16, color=COLOR_AZUL)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    _run(sub, f"Ref.: {ref}  ·  Matrícula: {mat}  ·  {marca} {model}",
         size=10, color=COLOR_GRIS)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(10)
    _run(sub2, f"Peticionario: {pet}  ·  Fecha: {hoy}",
         italic=True, size=9, color=COLOR_GRIS)

    _hr(doc, "1976D2")


# ════════════════════════════════════════════════════════════════════════════
#  Sección por bloque de cálculo
# ════════════════════════════════════════════════════════════════════════════

# Mapa de etiquetas legibles para claves de campos
ETIQUETAS = {
    # Geometría atornillada
    "TORN_TIPO_CORTE":     "Tipo de unión / esfuerzo",
    "TORN_NUM":            "Número de tornillos",
    "TORN_DIAM":           "Diámetro tornillo",
    "TORN_CALIDAD":        "Calidad / clase resistente",
    "TORN_A_BRUTA":        "Área bruta A (mm²)",
    "TORN_A_S":            "Área resistente As (mm²)",
    "TORN_FYB":            "Límite elástico fyb (MPa)",
    "TORN_FUB":            "Resistencia última fub (MPa)",
    "TORN_FV_RD_1":        "Fv,Rd por plano de corte (N)",
    "TORN_FV_RD":          "Fv,Rd total (N)",
    "TORN_FT_RD":          "Ft,Rd por tornillo (N)",
    "TORN_FV_ED":          "Fv,Ed — Cortante de cálculo (N)",
    "TORN_FT_ED":          "Ft,Ed — Tracción de cálculo (N)",
    "TORN_COEF_CORT":      "Coef. utilización cortante (Fv,Ed/Fv,Rd)",
    "TORN_COEF_TRAC":      "Coef. utilización tracción (Ft,Ed/Ft,Rd)",
    "TORN_INTER":          "Interacción cortante+tracción (≤1.0)",
    "TORN_RESULTADO":      "Resultado verificación",
    "TORN_NORMA":          "Norma de referencia",
    "TORN_NOTAS":          "Notas adicionales",
    # Exterior
    "TORN_EXT_CLASE_CORR": "Clase de corrosividad",
    "TORN_EXT_ACABADO":    "Acabado / protección tornillo",
    # Suspensión neumática
    "MARCA_BALONAS":       "Marca balonas",
    "PRESION_BALONAS_BAR": "Presión máx. balonas (bar)",
    "NUM_BALONAS":         "Número de balonas",
    "CARGA_TOTAL_KG":      "Carga total sobre eje (kg)",
    "CARGA_POR_BALONA_KG": "Carga por balona (kg)",
    "DIAM_BALONA_MM":      "Diámetro efectivo balona (mm)",
    "AREA_EFECTIVA_CM2":   "Área efectiva calculada (cm²)",
    "FUERZA_CALCULADA_N":  "Fuerza calculada (N)",
    "PRESION_TRABAJO_BAR": "Presión de trabajo (bar)",
    "MARGEN_SEGURIDAD":    "Margen de seguridad (%)",
    "NORMA_SUSPE":         "Norma / referencia",
    "NOTAS_SUSPE":         "Notas adicionales",
    # Motor / propulsión
    "MOTOR_MARCA":         "Marca del motor",
    "MOTOR_MODELO":        "Modelo / referencia",
    "MOTOR_CV":            "Potencia (CV)",
    "MOTOR_KW":            "Potencia (kW)",
    "MOTOR_PAR":           "Par máximo (Nm)",
    "MOTOR_RPM":           "Régimen par máximo (rpm)",
    "MOTOR_NORMA":         "Norma homologación",
    "MOTOR_NOTAS":         "Notas",
    # Suspensión neumática — campos calculados
    "SN_CARGA_BALONA":     "Carga por balona (kg)",
    "SN_AREA_EFECT":       "Área efectiva (cm²)",
    "SN_FUERZA_N":         "Fuerza calculada (N)",
    "SN_P_TRABAJO":        "Presión de trabajo (bar)",
    "SN_MARGEN":           "Margen de seguridad (%)",
    "SN_RESULTADO":        "Resultado verificación",
    # Suspensión mecánica
    "SM_VAR_CALC":         "Variación calculada (mm)",
    "SM_PCT_VAR":          "Variación porcentual (%)",
    # Cambio de motor
    "MOTOR_ORIG_TIPO":     "Motor original — Tipo",
    "MOTOR_ORIG_CILINDRADA_CC": "Motor original — Cilindrada (cc)",
    "MOTOR_ORIG_POT_KW":  "Motor original — Potencia (kW)",
    "MOTOR_ORIG_PAR_NM":  "Motor original — Par máx. (Nm)",
    "MOTOR_NUEVO_MARCA":  "Motor nuevo — Marca",
    "MOTOR_NUEVO_MODELO": "Motor nuevo — Modelo",
    "MOTOR_NUEVO_POT_KW": "Motor nuevo — Potencia (kW)",
    "MOTOR_NUEVO_PAR_NM": "Motor nuevo — Par máx. (Nm)",
    "MOTOR_TARA_KG":      "Tara del vehículo (kg)",
    "RATIO_POT_PESO":     "Relación potencia/peso (kW/t)",
    "INCREMENTO_POT_PCT": "Incremento de potencia (%)",
    "INCREMENTO_PAR_PCT": "Incremento de par (%)",
    "MOTOR_RESULTADO":    "Resultado verificación",
    # Conversión eléctrica
    "MOTOR_ELEC_POT_KW":  "Motor eléctrico — Potencia (kW)",
    "BATERIA_CAPACIDAD_KWH": "Batería — Capacidad (kWh)",
    "CONV_TARA_KG":       "Tara del vehículo (kg)",
    "CONV_RATIO_POT":     "Relación potencia/peso (kW/t)",
    "CONV_ENERGIA_PESO":  "Energía/peso (Wh/kg)",
    # Enganche
    "ENG_MMA_VEHICULO":   "MMA vehículo tractor (kg)",
    "ENG_MMA_CONJ_CALC":  "MMA conjunto calculada (kg)",
    "ENG_VERIF_CONJ":     "Verificación MMA conjunto",
    # Modificación de masas
    "MM_SUMA_EJES":       "Suma de cargas por eje (kg)",
    "MM_CARGA_UTIL":      "Carga útil (kg)",
    "MM_VERIF_MMA":       "Verificación MMA",
    "VERIFICACION_EJES":  "Verificación carga por eje",
    # Carrocería
    "CAR_VAR_LONG":       "Variación longitud (mm)",
    "CAR_VAR_ANCHO":      "Variación anchura (mm)",
    "CAR_VAR_ALTO":       "Variación altura (mm)",
    "CAR_VAR_VOL_DEL":    "Variación voladizo delantero (mm)",
    "CAR_VAR_VOL_TRA":    "Variación voladizo trasero (mm)",
    # Frenos modificación
    "FRMOD_MEJORA_PCT":   "Mejora distancia frenada (%)",
    "FRMOD_RESULTADO":    "Resultado verificación",
    # Flexión
    "FLEX_MATERIAL":      "Material estructural",
    "FLEX_FY_MPA":        "Límite elástico fy (MPa)",
    "FLEX_PERFIL":        "Referencia del perfil",
    "FLEX_W_CM3":         "Módulo resistente Wpl (cm³)",
    "FLEX_M_NM":          "Momento flector M (N·m)",
    "FLEX_SIGMA":         "σ_máx (MPa)",
    "FLEX_FD":            "fd (MPa)",
    "FLEX_COEF":          "Coef. utilización",
    "FLEX_RESULT":        "Resultado verificación",
    "FLEX_NORMA":         "Norma de referencia",
    # Torsión
    "TORS_TAU":           "τ_máx (MPa)",
    "TORS_TAU_ADM":       "τ_adm (MPa)",
    "TORS_COEF":          "Coef. utilización",
    "TORS_RESULT":        "Resultado verificación",
    "TORS_NORMA":         "Norma de referencia",
    # Pandeo
    "PAND_IRG":           "Radio de giro i (mm)",
    "PAND_ESBELT":        "Esbeltez λ",
    "PAND_PCR":           "Carga crítica Pcr (N)",
    "PAND_COEF":          "Coef. utilización",
    "PAND_RESULT":        "Resultado verificación",
    "PAND_NORMA":         "Norma de referencia",
    # Ejes
    "EJES_DQ1":           "Incremento eje delantero ΔQ₁ (kg)",
    "EJES_DQ2":           "Incremento eje trasero ΔQ₂ (kg)",
    "EJES_NQ1":           "Nueva carga eje delantero (kg)",
    "EJES_NQ2":           "Nueva carga eje trasero (kg)",
    "EJES_RES1":          "Verificación eje delantero",
    "EJES_RES2":          "Verificación eje trasero",
    # Vuelco
    "VUELT_ETA":          "Factor de estabilidad η",
    "VUELT_RESULT":       "Resultado verificación",
    "VUELT_NORMA":        "Norma de referencia",
    # Soldadura
    "SOLD_TAU":           "τ (MPa)",
    "SOLD_FW_RD":         "fw,Rd (MPa)",
    "SOLD_COEF":          "Coef. utilización",
    "SOLD_RESULT":        "Resultado verificación",
    "SOLD_NORMA":         "Norma de referencia",
    # UA Aerodinámica
    "UA_A_M2":            "Área de rozamiento (m²)",
    "UA_FX_N":            "Fuerza aerodinámica (N)",
    "UA_FC_N":            "Fuerza de cálculo (N)",
    "UA_NREQ":            "N.º tornillos necesarios",
    "UA_RESULT":          "Resultado verificación",
    "UA_NORMA":           "Norma de referencia",
    # Adhesiva aerodinámica
    "ADH_TAU":            "τ adhesivo (MPa)",
    "ADH_COEF":           "Coef. utilización",
    "ADH_RESULT":         "Resultado verificación",
    # Tacos
    "TACO_AR":            "Área resistente (mm²)",
    "TACO_SIGMA":         "σ compresión (MPa)",
    "TACO_COEF":          "Coef. utilización",
    "TACO_RESULT":        "Resultado verificación",
    # Protección trasera
    "PT_COEF":            "Coef. utilización",
    "PT_RESULT":          "Resultado verificación",
    # Frenos disco
    "FR_EFIC":            "Eficacia de frenado (%)",
    "FR_RESULT":          "Resultado verificación",
}

# Campos de resultado (se muestran con color según OK/NOK)
CAMPOS_RESULTADO = {
    "TORN_RESULTADO", "FLEX_RESULT", "TORS_RESULT", "PAND_RESULT",
    "EJES_RES1", "EJES_RES2", "VUELT_RESULT", "SOLD_RESULT",
    "UA_RESULT", "ADH_RESULT", "TACO_RESULT", "PT_RESULT", "FR_RESULT",
    "SN_RESULTADO", "FRMOD_RESULTADO", "MOTOR_RESULTADO",
    "ENG_VERIF_CONJ", "MM_VERIF_MMA", "VERIFICACION_EJES",
    "BM_VER1_SB", "BM_VER2_SB",
}
# Campos de comprobación numérica (se colorean según ≤1 o >1)
CAMPOS_COEF = {
    "TORN_COEF_CORT", "TORN_COEF_TRAC", "TORN_INTER",
    "FLEX_COEF", "TORS_COEF", "PAND_COEF", "SOLD_COEF",
    "ADH_COEF", "TACO_COEF", "PT_COEF",
}

# Separadores internos (claves que empiezan por _SEP_)
def _es_sep(clave):
    return clave.startswith("_SEP_")


# ════════════════════════════════════════════════════════════════════════════
#  Desarrollo paso a paso por tipo de cálculo
# ════════════════════════════════════════════════════════════════════════════

def _pasos_tornillos(doc, campos):
    """Uniones atornilladas — Interior / Exterior."""
    diam = _v(campos, "TORN_DIAM")
    cal  = _v(campos, "TORN_CALIDAD")
    n    = _v(campos, "TORN_NUM", "4")
    tipo_c = _v(campos, "TORN_TIPO_CORTE")
    A_s  = _v(campos, "TORN_A_S")
    A_b  = _v(campos, "TORN_A_BRUTA")
    fyb  = _v(campos, "TORN_FYB")
    fub  = _v(campos, "TORN_FUB")

    _paso(doc, 1, "Selección del tornillo y propiedades mecánicas")
    _descripcion(doc, f"Se selecciona tornillo métrico {diam}, clase resistente {cal}.")
    _descripcion(doc, f"Propiedades tabuladas: A = {A_b} mm², As = {A_s} mm², "
                      f"fyb = {fyb} MPa, fub = {fub} MPa. n = {n}. Unión: {tipo_c}.")

    Fv1  = _v(campos, "TORN_FV_RD_1")
    FvRd = _v(campos, "TORN_FV_RD")
    FtRd = _v(campos, "TORN_FT_RD")
    _paso(doc, 2, "Resistencias de cálculo (EN 1993-1-8, γM2 = 1,25)")
    _formula_omml(doc, [
        _omml_sub("F", "v,Rd"), " = ",
        _omml_frac([_omml_run("α"), _omml_sub("", "v"), _omml_run(" · "),
                     _omml_sub("f", "ub"), _omml_run(" · "),
                     _omml_sub("A", "s")],
                    [_omml_sub("γ", "M2")])
    ], f"Fv,Rd(1 plano) = {Fv1} N → Fv,Rd = {FvRd} N")
    _formula_omml(doc, [
        _omml_sub("F", "t,Rd"), " = ",
        _omml_frac([_omml_run("0,9 · "), _omml_sub("f", "ub"),
                     _omml_run(" · "), _omml_sub("A", "s")],
                    [_omml_sub("γ", "M2")])
    ], f"Ft,Rd = {FtRd} N")

    FvEd = _v(campos, "TORN_FV_ED")
    FtEd = _v(campos, "TORN_FT_ED", "0")
    _paso(doc, 3, "Cargas aplicadas")
    _descripcion(doc, f"Fv,Ed = {FvEd} N   ·   Ft,Ed = {FtEd} N")

    cc = _v(campos, "TORN_COEF_CORT")
    ct = _v(campos, "TORN_COEF_TRAC")
    inter = _v(campos, "TORN_INTER")
    _paso(doc, 4, "Comprobaciones de seguridad")
    _formula_omml(doc, [
        _omml_frac([_omml_sub("F", "v,Ed")], [_omml_sub("F", "v,Rd")]),
        " ≤ 1,0"
    ], f"= {FvEd} / {FvRd} = {cc}")
    _formula_omml(doc, [
        _omml_frac([_omml_sub("F", "v,Ed")], [_omml_sub("F", "v,Rd")]),
        " + ",
        _omml_frac([_omml_sub("F", "t,Ed")],
                    [_omml_run("1,4 · "), _omml_sub("F", "t,Rd")]),
        " ≤ 1,0"
    ], f"Interacción = {inter}")

    res = _v(campos, "TORN_RESULTADO")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_flexion(doc, campos):
    """Flexión simple de viga — Navier."""
    mat   = _v(campos, "FLEX_MATERIAL")
    perfil = _v(campos, "FLEX_PERFIL")
    fy    = _v(campos, "FLEX_FY_MPA")
    W     = _v(campos, "FLEX_W_CM3")
    M     = _v(campos, "FLEX_M_NM")
    sigma = _v(campos, "FLEX_SIGMA")
    fd    = _v(campos, "FLEX_FD")
    coef  = _v(campos, "FLEX_COEF")

    _paso(doc, 1, "Datos de la sección y material")
    _descripcion(doc, f"Material: {mat} (fy = {fy} MPa). Perfil: {perfil}. Wpl = {W} cm³.")

    _paso(doc, 2, "Tensión máxima por flexión (Navier)")
    _formula_omml(doc, [
        _omml_sub("σ", "máx"), " = ",
        _omml_frac("M", [_omml_sub("W", "pl")])
    ], f"σ_máx = {M} / {W} = {sigma} MPa")

    _paso(doc, 3, "Resistencia de cálculo (γM0 = 1,0)")
    _formula_omml(doc, [
        _omml_sub("f", "d"), " = ",
        _omml_frac([_omml_sub("f", "y")], [_omml_sub("γ", "M0")])
    ], f"fd = {fy} / 1,0 = {fd} MPa")

    _paso(doc, 4, "Verificación")
    _formula_omml(doc, [
        _omml_frac([_omml_sub("σ", "máx")], [_omml_sub("f", "d")]),
        " ≤ 1,0"
    ], f"= {sigma} / {fd} = {coef}")
    res = _v(campos, "FLEX_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_torsion(doc, campos):
    """Torsión de perfil — Saint-Venant."""
    r   = _v(campos, "TORS_R_MM")
    Ip  = _v(campos, "TORS_IP_CM4")
    fy  = _v(campos, "TORS_FY_MPA")
    MT  = _v(campos, "TORS_MT_NM")
    tau = _v(campos, "TORS_TAU")
    tau_adm = _v(campos, "TORS_TAU_ADM")
    coef = _v(campos, "TORS_COEF")

    _paso(doc, 1, "Datos de la sección")
    _descripcion(doc, f"Tipo: {_v(campos, 'TORS_TIPO')}. r = {r} mm. Ip = {Ip} cm⁴. fy = {fy} MPa.")

    _paso(doc, 2, "Tensión tangencial máxima")
    _formula_omml(doc, [
        _omml_sub("τ", "máx"), " = ",
        _omml_frac([_omml_sub("M", "T"), _omml_run(" · r")],
                    [_omml_sub("I", "p")])
    ], f"τ = {MT} · {r} / {Ip} = {tau} MPa")

    _paso(doc, 3, "Tensión admisible (Von Mises)")
    _formula_omml(doc, [
        _omml_sub("τ", "adm"), " = ",
        _omml_frac([_omml_sub("f", "y")],
                    [_omml_sqrt("3")])
    ], f"τ_adm = {fy} / √3 = {tau_adm} MPa")

    _paso(doc, 4, "Verificación")
    _formula_omml(doc, [
        _omml_frac([_omml_sub("τ", "máx")], [_omml_sub("τ", "adm")]),
        " ≤ 1,0"
    ], f"= {tau} / {tau_adm} = {coef}")
    res = _v(campos, "TORS_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_pandeo(doc, campos):
    """Pandeo de barra — Euler."""
    mat  = _v(campos, "PAND_MATERIAL")
    fy   = _v(campos, "PAND_FY_MPA")
    E    = _v(campos, "PAND_E_MPA", "210000")
    L    = _v(campos, "PAND_L_MM")
    I    = _v(campos, "PAND_I_CM4")
    A    = _v(campos, "PAND_A_CM2")
    mu   = _v(campos, "PAND_MU")
    irg  = _v(campos, "PAND_IRG")
    esb  = _v(campos, "PAND_ESBELT")
    Pcr  = _v(campos, "PAND_PCR")
    N    = _v(campos, "PAND_N_N")
    coef = _v(campos, "PAND_COEF")

    _paso(doc, 1, "Datos de la barra")
    _descripcion(doc, f"Material: {mat} (fy = {fy} MPa, E = {E} MPa). "
                      f"L = {L} mm. Imin = {I} cm⁴. A = {A} cm². μ = {mu}.")

    _paso(doc, 2, "Radio de giro y esbeltez")
    _formula_omml(doc, [
        "i = ",
        _omml_sqrt([_omml_frac("I", "A")])
    ], f"i = {irg} mm")
    _formula_omml(doc, [
        "λ = ",
        _omml_frac(["μ · L"], ["i"])
    ], f"λ = {mu} · {L} / {irg} = {esb}")

    _paso(doc, 3, "Carga crítica de Euler")
    _formula_omml(doc, [
        _omml_sub("P", "cr"), " = ",
        _omml_frac([_omml_sup("π", "2"), _omml_run(" · E · I")],
                    [_omml_sup("(μ · L)", "2")])
    ], f"Pcr = {Pcr} N")

    _paso(doc, 4, "Verificación")
    _formula_omml(doc, [
        _omml_frac("N", [_omml_sub("P", "cr")]),
        " ≤ 1,0"
    ], f"= {N} / {Pcr} = {coef}")
    res = _v(campos, "PAND_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_ejes(doc, campos):
    """Distribución de cargas por ejes."""
    L  = _v(campos, "EJES_BATALL")
    P  = _v(campos, "EJES_P_KG")
    d  = _v(campos, "EJES_D_MM")
    Q1 = _v(campos, "EJES_Q1_KG")
    Q2 = _v(campos, "EJES_Q2_KG")

    _paso(doc, 1, "Datos del vehículo y carga")
    _descripcion(doc, f"Batalla: L = {L} mm. Peso carga: P = {P} kg.")
    _descripcion(doc, f"Distancia CG carga a eje delantero: d = {d} mm.")
    _descripcion(doc, f"Cargas actuales: Q₁ = {Q1} kg (del.), Q₂ = {Q2} kg (tras.).")

    _paso(doc, 2, "Incremento de carga por eje (equilibrio estático)")
    dQ1 = _v(campos, "EJES_DQ1")
    dQ2 = _v(campos, "EJES_DQ2")
    _formula(doc, "ΔQ₂ = P · d / L     (incremento eje trasero)",
                  f"ΔQ₂ = {P} · {d} / {L} = {dQ2} kg")
    _formula(doc, "ΔQ₁ = P − ΔQ₂       (incremento eje delantero)",
                  f"ΔQ₁ = {P} − {dQ2} = {dQ1} kg")

    _paso(doc, 3, "Nuevas cargas por eje")
    NQ1 = _v(campos, "EJES_NQ1")
    NQ2 = _v(campos, "EJES_NQ2")
    _formula(doc, "Q₁_nuevo = Q₁ + ΔQ₁",
                  f"Q₁_nuevo = {Q1} + {dQ1} = {NQ1} kg")
    _formula(doc, "Q₂_nuevo = Q₂ + ΔQ₂",
                  f"Q₂_nuevo = {Q2} + {dQ2} = {NQ2} kg")

    _paso(doc, 4, "Verificación frente a MMA por eje")
    MMA1 = _v(campos, "EJES_MMA1")
    MMA2 = _v(campos, "EJES_MMA2")
    r1   = _v(campos, "EJES_RES1")
    r2   = _v(campos, "EJES_RES2")
    _descripcion(doc, f"MMA eje delantero: {MMA1} kg → Q₁_nuevo = {NQ1} kg")
    _resultado_paso(doc, f"Eje delantero: {r1}", "✓" in r1)
    _descripcion(doc, f"MMA eje trasero: {MMA2} kg → Q₂_nuevo = {NQ2} kg")
    _resultado_paso(doc, f"Eje trasero: {r2}", "✓" in r2)


def _pasos_vuelco(doc, campos):
    """Estabilidad al vuelco lateral."""
    s   = _v(campos, "VUELT_ANCHO")
    hcg = _v(campos, "VUELT_HCG")
    eta  = _v(campos, "VUELT_ETA")
    alim = _v(campos, "VUELT_ALIM")
    vlim = _v(campos, "VUELT_VLIM")

    _paso(doc, 1, "Datos geométricos del vehículo")
    _descripcion(doc, f"Ancho de vía: s = {s} mm. Altura del CG: hcg = {hcg} mm.")

    _paso(doc, 2, "Factor de estabilidad η")
    _formula_omml(doc, [
        "η = ",
        _omml_frac("s", [_omml_run("2 · "), _omml_sub("h", "cg")])
    ], f"η = {s} / (2 · {hcg}) = {eta}")

    _paso(doc, 3, "Aceleración lateral y velocidad límite")
    _formula_omml(doc, [
        _omml_sub("a", "lat"), " = η · g"
    ], f"a_lat = {eta} · 9,81 = {alim} m/s²")
    _descripcion(doc, f"Velocidad límite en curva R = 50 m: V = {vlim} km/h.")

    _paso(doc, 4, "Verificación (η ≥ 0,3 según RD 2028/86)")
    res = _v(campos, "VUELT_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res or "ESTABLE" in res)


def _pasos_soldadura(doc, campos):
    """Uniones soldadas — cordón en ángulo."""
    mat  = _v(campos, "SOLD_MATERIAL")
    fu   = _v(campos, "SOLD_FU_MPA")
    beta = _v(campos, "SOLD_BETA")
    a    = _v(campos, "SOLD_GARG")
    Lw   = _v(campos, "SOLD_LONG")
    F    = _v(campos, "SOLD_F_N")
    tau  = _v(campos, "SOLD_TAU")
    fw   = _v(campos, "SOLD_FW_RD")
    coef = _v(campos, "SOLD_COEF")

    _paso(doc, 1, "Datos del material y la soldadura")
    _descripcion(doc, f"Material: {mat}. fu = {fu} MPa, βw = {beta}. "
                      f"Garganta: a = {a} mm. Cordón: Lw = {Lw} mm.")

    _paso(doc, 2, "Tensión tangencial en el cordón")
    _formula_omml(doc, [
        "τ = ",
        _omml_frac("F", [_omml_run("a · "), _omml_sub("L", "w")])
    ], f"τ = {F} / ({a} · {Lw}) = {tau} MPa")

    _paso(doc, 3, "Resistencia del cordón (EN 1993-1-8 §4.5.3.3, γM2 = 1,25)")
    _formula_omml(doc, [
        _omml_sub("f", "w,Rd"), " = ",
        _omml_frac([_omml_sub("f", "u")],
                    [_omml_sqrt("3"), _omml_run(" · "),
                     _omml_sub("β", "w"), _omml_run(" · "),
                     _omml_sub("γ", "M2")])
    ], f"fw,Rd = {fu} / (√3 · {beta} · 1,25) = {fw} MPa")

    _paso(doc, 4, "Verificación")
    _formula_omml(doc, [
        _omml_frac("τ", [_omml_sub("f", "w,Rd")]),
        " ≤ 1,0"
    ], f"= {tau} / {fw} = {coef}")
    res = _v(campos, "SOLD_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_grua(doc, campos):
    """Grúa autocarga — momento de vuelco."""
    P    = _v(campos, "GRUA_CAP_KG")
    L    = _v(campos, "GRUA_ALC_MM")
    bcil = _v(campos, "GRUA_BCIL_MM")

    _paso(doc, 1, "Datos de la grúa")
    _descripcion(doc, f"Capacidad: P = {P} kg. Alcance: L = {L} mm. "
                      f"Brazo del cilindro: b_cil = {bcil} mm.")

    _paso(doc, 2, "Momento de vuelco")
    Mv = _v(campos, "GRUA_MV_NM")
    _formula(doc, "Mv = P · g · L",
                  f"Mv = {P} · 9,81 · {L} = {Mv} N·m")

    _paso(doc, 3, "Fuerza en el cilindro")
    Fcil = _v(campos, "GRUA_FCIL_N")
    _formula(doc, "F_cil = Mv / b_cil",
                  f"F_cil = {Mv} · 1000 / {bcil} = {Fcil} N")


def _pasos_balance_masas(doc, campos):
    """Balance de masas — Turismo (M1/M2)."""
    MMA   = _v(campos, "BM_MMA")
    MMAe1 = _v(campos, "BM_MMA_E1")
    MMAe2 = _v(campos, "BM_MMA_E2")
    mT1   = _v(campos, "BM_MT1")
    mT2   = _v(campos, "BM_MT2")

    _paso(doc, 1, "Datos del vehículo")
    _descripcion(doc, f"MMA = {MMA} kg. MMA eje 1 = {MMAe1} kg. MMA eje 2 = {MMAe2} kg.")
    _descripcion(doc, f"Tara eje 1 = {mT1} kg. Tara eje 2 = {mT2} kg.")

    nP1 = _v(campos, "BM_NP1", "2")
    nP2 = _v(campos, "BM_NP2", "0")
    _paso(doc, 2, "Ocupantes")
    _descripcion(doc, f"Fila 1: {nP1} pasajeros × 75 kg. Fila 2: {nP2} pasajeros × 75 kg.")

    _paso(doc, 3, "Resultados — Reacciones en ejes")
    tara = _v(campos, "BM_TARA")
    mQu  = _v(campos, "BM_MQU")
    R1   = _v(campos, "BM_R1_SB")
    R2   = _v(campos, "BM_R2_SB")
    _descripcion(doc, f"Tara total = {tara} kg. Carga útil mQu = {mQu} kg.")
    _formula(doc, "R₁ y R₂ se calculan por equilibrio estático de momentos",
                  f"R₁ = {R1} kgf, R₂ = {R2} kgf")

    _paso(doc, 4, "Verificación frente a MMA por eje")
    v1 = _v(campos, "BM_VER1_SB")
    v2 = _v(campos, "BM_VER2_SB")
    _resultado_paso(doc, f"Eje delantero: {v1}", "✓" in v1)
    _resultado_paso(doc, f"Eje trasero: {v2}", "✓" in v2)


def _pasos_ua_aero(doc, campos):
    """Unión atornillada — Carga aerodinámica."""
    Cx = _v(campos, "UA_CX")
    V  = _v(campos, "UA_V_KMH")
    w  = _v(campos, "UA_ANCHO")
    h  = _v(campos, "UA_ALTO")
    Pp = _v(campos, "UA_PP_KG")

    _paso(doc, 1, "Superficie expuesta y carga aerodinámica")
    A  = _v(campos, "UA_A_M2")
    Fx = _v(campos, "UA_FX_N")
    Fc = _v(campos, "UA_FC_N")
    _formula(doc, "A = ancho · alto / 10.000",
                  f"A = {w} · {h} / 10.000 = {A} m²")
    _formula(doc, "Fx = 0,5 · ρ · V² · Cx · A     [ρ = 1,225 kg/m³]",
                  f"Fx = {Fx} N")
    _formula(doc, "F_calc = Fx + Pp · g",
                  f"F_calc = {Fx} + {Pp} · 9,81 = {Fc} N")

    _paso(doc, 2, "Selección del tornillo")
    met = _v(campos, "UA_METRICA")
    cal = _v(campos, "UA_CALIDAD")
    cha = _v(campos, "UA_CHAPA")
    _descripcion(doc, f"Métrica: {met}. Calidad: {cal}. Chapa: {cha}.")

    _paso(doc, 3, "Número de tornillos necesarios")
    Nt_t = _v(campos, "UA_NT_TRAC")
    Nt_c = _v(campos, "UA_NT_CORT")
    Nt_a = _v(campos, "UA_NT_APC")
    Nreq = _v(campos, "UA_NREQ")
    _formula(doc, "Nt (tracción, cortadura, aplastamiento) → máximo",
                  f"Nt_trac = {Nt_t}, Nt_cort = {Nt_c}, Nt_aplast = {Nt_a}")
    _descripcion(doc, f"Número de tornillos necesarios: {Nreq}")

    res = _v(campos, "UA_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_adh_aero(doc, campos):
    """Unión adhesiva — Carga aerodinámica."""
    Cx = _v(campos, "ADH_CX")
    V  = _v(campos, "ADH_V_KMH")
    Pp = _v(campos, "ADH_PP_KG")

    _paso(doc, 1, "Carga aerodinámica")
    A  = _v(campos, "ADH_A_M2")
    Fx = _v(campos, "ADH_FX_N")
    Fc = _v(campos, "ADH_FC_N")
    _formula(doc, "Fx = 0,5 · ρ · V² · Cx · A",
                  f"Fx = {Fx} N. F_calc = {Fc} N")

    _paso(doc, 2, "Datos del adhesivo")
    adh  = _v(campos, "ADH_TIPO")
    R    = _v(campos, "ADH_R_MPA")
    b    = _v(campos, "ADH_B_MM")
    l    = _v(campos, "ADH_L_MM")
    _descripcion(doc, f"Adhesivo: {adh}. Resistencia a cortante: R = {R} MPa.")
    _descripcion(doc, f"Área de unión: b = {b} mm × l = {l} mm.")

    _paso(doc, 3, "Verificación a cortante")
    tau  = _v(campos, "ADH_TAU")
    coef = _v(campos, "ADH_COEF")
    _formula(doc, "τ = F_calc / (b · l)",
                  f"τ = {Fc} / ({b} · {l}) = {tau} MPa")
    _formula(doc, "Coef. = τ / R ≤ 1,0",
                  f"Coef. = {tau} / {R} = {coef}")
    res = _v(campos, "ADH_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_tacos(doc, campos):
    """Tacos de elevación — Suspensión / Altura."""
    mat  = _v(campos, "TACO_MAT")
    MTMA = _v(campos, "TACO_MTMA")
    n    = _v(campos, "TACO_N")
    d    = _v(campos, "TACO_D_MM")
    Rc   = _v(campos, "TACO_RCOMP")
    Ar   = _v(campos, "TACO_AR")
    sigma = _v(campos, "TACO_SIGMA")
    coef = _v(campos, "TACO_COEF")

    _paso(doc, 1, "Datos de los tacos")
    _descripcion(doc, f"Material: {mat}. R_comp = {Rc} MPa. MTMA = {MTMA} kg. n = {n}. d = {d} mm.")

    _paso(doc, 2, "Área resistente de cada taco")
    _formula_omml(doc, [
        _omml_sub("A", "r"), " = ",
        _omml_frac([_omml_run("π · "), _omml_sup("d", "2")], ["4"])
    ], f"Ar = π · {d}² / 4 = {Ar} mm²")

    _paso(doc, 3, "Tensión de compresión")
    _formula_omml(doc, [
        "σ = ",
        _omml_frac(["MTMA · g"], [_omml_run("n · "), _omml_sub("A", "r")])
    ], f"σ = ({MTMA} · 9,81) / ({n} · {Ar}) = {sigma} MPa")

    _paso(doc, 4, "Verificación")
    _formula_omml(doc, [
        _omml_frac("σ", [_omml_sub("R", "comp")]),
        " ≤ 1,0"
    ], f"= {sigma} / {Rc} = {coef}")
    res = _v(campos, "TACO_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_proteccion_trasera(doc, campos):
    """Protección trasera — barra antichoque."""
    MTMA = _v(campos, "PT_MTMA")
    L    = _v(campos, "PT_LONG")
    W    = _v(campos, "PT_W_CM3")
    mat  = _v(campos, "PT_MAT")

    _paso(doc, 1, "Datos de la barra de protección")
    _descripcion(doc, f"MTMA = {MTMA} kg. Longitud libre: L = {L} mm.")
    _descripcion(doc, f"Módulo resistente: W = {W} cm³. Material: {mat}.")

    _paso(doc, 2, "Fuerza de cálculo (50% carga en eje trasero)")
    Fh = _v(campos, "PT_F_HALF")
    _formula(doc, "F = 0,5 · MTMA · g / 2",
                  f"F = {Fh} N")

    _paso(doc, 3, "Momento flector de cálculo")
    Med = _v(campos, "PT_M_ED")
    _formula(doc, "M_Ed = F · L / 4     (viga biapoyada, carga centrada)",
                  f"M_Ed = {Med} N·mm")

    _paso(doc, 4, "Verificación a flexión")
    sadm = _v(campos, "PT_SADM")
    Mrd  = _v(campos, "PT_M_RD")
    coef = _v(campos, "PT_COEF")
    _formula(doc, "M_Rd = σ_adm · W",
                  f"σ_adm = {sadm} MPa → M_Rd = {Mrd} N·mm")
    _formula(doc, "Coef. = M_Ed / M_Rd ≤ 1,0",
                  f"Coef. = {Med} / {Mrd} = {coef}")
    res = _v(campos, "PT_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_frenos_disco(doc, campos):
    """Frenos — disco (comprobación eficacia)."""
    MMA = _v(campos, "FR_MMA")
    mu  = _v(campos, "FR_MU")

    _paso(doc, 1, "Datos del vehículo y neumáticos")
    asp = _v(campos, "FR_ASPECTO")
    sec = _v(campos, "FR_SECCION")
    lla = _v(campos, "FR_LLANTA")
    R   = _v(campos, "FR_RLLANTA")
    _descripcion(doc, f"MMA = {MMA} kg. Coef. rozamiento μ = {mu}.")
    _descripcion(doc, f"Neumático: {sec}/{asp} R{lla}. Radio dinámico: R = {R} m.")

    _paso(doc, 2, "Par de frenado por eje")
    Td = _v(campos, "FR_T_DEL")
    Tt = _v(campos, "FR_T_TRA")
    Ttot = _v(campos, "FR_T_TOT")
    _descripcion(doc, f"Par frenado delantero: T_del = {Td} N·m.")
    _descripcion(doc, f"Par frenado trasero: T_tra = {Tt} N·m.")
    _formula(doc, "T_total = T_del + T_tra",
                  f"T_total = {Td} + {Tt} = {Ttot} N·m")

    _paso(doc, 3, "Fuerza de frenado y eficacia")
    Ff   = _v(campos, "FR_F_FREN")
    efic = _v(campos, "FR_EFIC")
    _formula(doc, "F_fren = T_total / R",
                  f"F_fren = {Ttot} / {R} = {Ff} N")
    _formula(doc, "Eficacia = F_fren / (MMA · g) · 100",
                  f"Eficacia = {efic} %")

    _paso(doc, 4, "Verificación (eficacia ≥ 50%)")
    res = _v(campos, "FR_RESULT")
    _resultado_paso(doc, f"→ {res}", "✓" in res)


def _pasos_suspension_neumatica(doc, campos):
    """Suspensión neumática — Balonas / Airbag."""
    marca = _v(campos, "MARCA_BALONAS")
    pmax  = _v(campos, "PRESION_BALONAS_BAR")
    nbal  = _v(campos, "NUM_BALONAS")
    diam  = _v(campos, "DIAM_BALONA_MM")
    carga = _v(campos, "CARGA_TOTAL_KG")
    cb    = _v(campos, "SN_CARGA_BALONA")
    area  = _v(campos, "SN_AREA_EFECT")
    F     = _v(campos, "SN_FUERZA_N")
    pt    = _v(campos, "SN_P_TRABAJO")
    margen = _v(campos, "SN_MARGEN")

    _paso(doc, 1, "Datos del sistema neumático")
    _descripcion(doc, f"Marca: {marca}. Presión máxima: {pmax} bar. "
                      f"Balonas: {nbal} × Ø{diam} mm. Carga eje: {carga} kg.")

    _paso(doc, 2, "Carga por balona")
    _formula_omml(doc, [
        _omml_sub("q", "bal"), " = ",
        _omml_frac([_omml_sub("Q", "total")], [_omml_sub("N", "bal")])
    ], f"q_bal = {carga} / {nbal} = {cb} kg")

    _paso(doc, 3, "Área efectiva y fuerza")
    _formula_omml(doc, [
        _omml_sub("A", "ef"), " = ",
        _omml_frac([_omml_run("π · "), _omml_sup("d", "2")],
                    ["4 · 100"])
    ], f"A_ef = π · {diam}² / 400 = {area} cm²")
    _formula_omml(doc, [
        "F = ", _omml_sub("Q", "total"), " · g"
    ], f"F = {carga} · 9,81 = {F} N")

    _paso(doc, 4, "Presión de trabajo y margen de seguridad")
    _formula_omml(doc, [
        _omml_sub("P", "trabajo"), " = ",
        _omml_frac("F", [_omml_sub("A", "ef"), _omml_run(" · 100.000")])
    ], f"P_trabajo = {F} / ({area} · 100.000) = {pt} bar")
    _formula_omml(doc, [
        "Margen = ",
        _omml_frac([_omml_sub("P", "máx")], [_omml_sub("P", "trabajo")]),
        " − 1) · 100"
    ], f"Margen = ({pmax} / {pt} − 1) · 100 = {margen} %")

    res = _v(campos, "SN_RESULTADO")
    _resultado_paso(doc, f"→ {res}", "✓" in str(res))


def _pasos_suspension_mecanica(doc, campos):
    """Suspensión mecánica — Sustitución muelles / amortiguadores."""
    _paso(doc, 1, "Datos de la sustitución")
    marca_m = _v(campos, "MARCA_MUELLE_NUEVO")
    marca_a = _v(campos, "MARCA_AMORT_NUEVO")
    antes   = _v(campos, "ALTURA_ANTES_MM")
    despues = _v(campos, "ALTURA_DESPUES_MM")
    _descripcion(doc, f"Muelles nuevos: {marca_m}. Amortiguadores nuevos: {marca_a}.")
    _descripcion(doc, f"Altura vehículo antes: {antes} mm. Después: {despues} mm.")

    _paso(doc, 2, "Variación de altura")
    var = _v(campos, "SM_VAR_CALC")
    pct = _v(campos, "SM_PCT_VAR")
    _formula(doc, "Variación = Altura_después − Altura_antes",
                  f"Variación = {despues} − {antes} = {var} mm")
    _formula(doc, "Variación (%) = Variación / Altura_antes · 100",
                  f"Variación = {var} / {antes} · 100 = {pct} %")


def _pasos_motor(doc, campos):
    """Cambio de motor / Sustitución unidad motriz."""
    _paso(doc, 1, "Motor original")
    _descripcion(doc, f"Tipo: {_v(campos, 'MOTOR_ORIG_TIPO')}. "
                      f"Cilindrada: {_v(campos, 'MOTOR_ORIG_CILINDRADA_CC')} cc. "
                      f"Potencia: {_v(campos, 'MOTOR_ORIG_POT_KW')} kW. "
                      f"Par: {_v(campos, 'MOTOR_ORIG_PAR_NM')} Nm.")

    _paso(doc, 2, "Motor nuevo")
    _descripcion(doc, f"Marca: {_v(campos, 'MOTOR_NUEVO_MARCA')}. "
                      f"Modelo: {_v(campos, 'MOTOR_NUEVO_MODELO')}.")
    _descripcion(doc, f"Cilindrada: {_v(campos, 'MOTOR_NUEVO_CILINDRADA_CC')} cc. "
                      f"Potencia: {_v(campos, 'MOTOR_NUEVO_POT_KW')} kW. "
                      f"Par: {_v(campos, 'MOTOR_NUEVO_PAR_NM')} Nm.")

    _paso(doc, 3, "Cálculos comparativos")
    ratio   = _v(campos, "RATIO_POT_PESO")
    inc_pot = _v(campos, "INCREMENTO_POT_PCT")
    inc_par = _v(campos, "INCREMENTO_PAR_PCT")
    _formula(doc, "Ratio potencia/peso = Pot_nueva / (Tara / 1000)",
                  f"Ratio = {ratio} kW/t")
    _formula(doc, "Incremento potencia = (Pot_nueva − Pot_orig) / Pot_orig · 100",
                  f"Incremento potencia = {inc_pot} %")
    _formula(doc, "Incremento par = (Par_nuevo − Par_orig) / Par_orig · 100",
                  f"Incremento par = {inc_par} %")


def _pasos_conversion(doc, campos):
    """Conversión eléctrica o híbrida."""
    _paso(doc, 1, "Motor eléctrico")
    _descripcion(doc, f"Marca: {_v(campos, 'MOTOR_ELEC_MARCA')}. "
                      f"Modelo: {_v(campos, 'MOTOR_ELEC_MODELO')}. "
                      f"Potencia: {_v(campos, 'MOTOR_ELEC_POT_KW')} kW.")

    _paso(doc, 2, "Batería")
    _descripcion(doc, f"Marca: {_v(campos, 'BATERIA_MARCA')}. "
                      f"Capacidad: {_v(campos, 'BATERIA_CAPACIDAD_KWH')} kWh. "
                      f"Tensión: {_v(campos, 'BATERIA_TENSION_V')} V.")

    _paso(doc, 3, "Ratios calculados")
    ratio = _v(campos, "CONV_RATIO_POT")
    ep    = _v(campos, "CONV_ENERGIA_PESO")
    _formula(doc, "Ratio potencia/peso = Pot / (Tara / 1000)",
                  f"Ratio = {ratio} kW/t")
    _formula(doc, "Energía/peso = Capacidad · 1000 / Tara",
                  f"Energía/peso = {ep} Wh/kg")


def _pasos_enganche(doc, campos):
    """Enganche / Dispositivo de acoplamiento."""
    _paso(doc, 1, "Datos del enganche")
    _descripcion(doc, f"Marca: {_v(campos, 'ENGANCHE_MARCA')}. "
                      f"Tipo: {_v(campos, 'ENGANCHE_TIPO')}. "
                      f"Homologación: {_v(campos, 'ENGANCHE_HOMOLOG')}.")

    _paso(doc, 2, "Masas remolcables")
    _descripcion(doc, f"MMA vehículo tractor: {_v(campos, 'ENG_MMA_VEHICULO')} kg.")
    _descripcion(doc, f"Masa máx. remolcable: {_v(campos, 'MASA_REMOLCABLE_KG')} kg.")

    _paso(doc, 3, "Verificación MMA del conjunto")
    mma_calc = _v(campos, "ENG_MMA_CONJ_CALC")
    mma_conj = _v(campos, "MMA_CONJUNTO_KG")
    _formula(doc, "MMA_conjunto_calc = MMA_tractor + Masa_remolcable",
                  f"MMA_conjunto_calc = {mma_calc} kg")
    _descripcion(doc, f"MMA conjunto declarada: {mma_conj} kg.")
    res = _v(campos, "ENG_VERIF_CONJ")
    _resultado_paso(doc, f"→ {res}", "✓" in str(res))


def _pasos_masas(doc, campos):
    """Modificación de masas / MMA."""
    _paso(doc, 1, "Datos de masas")
    _descripcion(doc, f"MMA anterior: {_v(campos, 'MMA_ANTERIOR_KG')} kg. "
                      f"MMA nueva: {_v(campos, 'MMA_NUEVA_KG')} kg.")
    _descripcion(doc, f"Tara: {_v(campos, 'TARA_KG')} kg. "
                      f"MMA máxima legal: {_v(campos, 'MMA_MAX_LEGAL_KG')} kg.")

    _paso(doc, 2, "Carga útil")
    cu = _v(campos, "MM_CARGA_UTIL")
    _formula(doc, "Carga útil = MMA − Tara",
                  f"Carga útil = {cu} kg")

    _paso(doc, 3, "Verificación de cargas por eje")
    suma = _v(campos, "MM_SUMA_EJES")
    _formula(doc, "Suma ejes = Eje1 + Eje2 + Eje3",
                  f"Suma ejes = {suma} kg")

    v_mma = _v(campos, "MM_VERIF_MMA")
    v_eje = _v(campos, "VERIFICACION_EJES")
    _resultado_paso(doc, f"MMA: {v_mma}", "✓" in str(v_mma))
    _resultado_paso(doc, f"Ejes: {v_eje}", "✓" in str(v_eje))


def _pasos_carroceria(doc, campos):
    """Modificación de carrocería / Estructura."""
    _paso(doc, 1, "Dimensiones del vehículo")
    _descripcion(doc, f"Material carrocería: {_v(campos, 'MATERIAL_CARROCERIA')}.")

    _paso(doc, 2, "Variaciones dimensionales calculadas")
    pares = [
        ("Longitud",           "CAR_VAR_LONG"),
        ("Anchura",            "CAR_VAR_ANCHO"),
        ("Altura",             "CAR_VAR_ALTO"),
        ("Voladizo delantero", "CAR_VAR_VOL_DEL"),
        ("Voladizo trasero",   "CAR_VAR_VOL_TRA"),
    ]
    for nombre, clave in pares:
        val = _v(campos, clave)
        if val and val != "—":
            _descripcion(doc, f"Variación {nombre}: {val} mm")


def _pasos_frenos_mod(doc, campos):
    """Modificación del sistema de frenos."""
    _paso(doc, 1, "Datos del sistema de frenos")
    _descripcion(doc, f"Tipo: {_v(campos, 'FRENO_SISTEMA_TIPO')}. "
                      f"Componente: {_v(campos, 'FRENO_MARCA_COMP')} {_v(campos, 'FRENO_MODELO_COMP')}.")

    _paso(doc, 2, "Comparación de distancia de frenada")
    antes = _v(campos, "DIST_FRENADA_ANTES")
    desp  = _v(campos, "DIST_FRENADA_DESP")
    mejora = _v(campos, "FRMOD_MEJORA_PCT")
    _formula(doc, "Mejora = (D_antes − D_después) / D_antes · 100",
                  f"Mejora = ({antes} − {desp}) / {antes} · 100 = {mejora} %")
    res = _v(campos, "FRMOD_RESULTADO")
    _resultado_paso(doc, f"→ {res}", "✓" in str(res))


# ── Dispatcher ──────────────────────────────────────────────────────────────

_PASOS_MAP = {
    "Uniones atornilladas — Interior":                     _pasos_tornillos,
    "Uniones atornilladas — Exterior (ambiente agresivo)":  _pasos_tornillos,
    "Flexión simple de viga — Navier":                      _pasos_flexion,
    "Torsión de perfil — Saint-Venant":                     _pasos_torsion,
    "Pandeo de barra — Euler":                              _pasos_pandeo,
    "Distribución de cargas por ejes":                      _pasos_ejes,
    "Estabilidad al vuelco lateral":                        _pasos_vuelco,
    "Uniones soldadas — cordón en ángulo":                  _pasos_soldadura,
    "Grúa autocarga — momento de vuelco":                   _pasos_grua,
    "Balance de masas — Turismo (M1/M2)":                   _pasos_balance_masas,
    "Unión atornillada — Carga aerodinámica":               _pasos_ua_aero,
    "Unión adhesiva — Carga aerodinámica":                  _pasos_adh_aero,
    "Tacos de elevación — Suspensión / Altura":             _pasos_tacos,
    "Protección trasera — barra antichoque":                _pasos_proteccion_trasera,
    "Frenos — disco (comprobación eficacia)":               _pasos_frenos_disco,
    "Suspensión neumática — Balonas / Airbag":              _pasos_suspension_neumatica,
    "Suspensión mecánica — Sustitución muelles / amortiguadores": _pasos_suspension_mecanica,
    "Cambio de motor / Sustitución unidad motriz":          _pasos_motor,
    "Conversión eléctrica o híbrida":                       _pasos_conversion,
    "Enganche / Dispositivo de acoplamiento":               _pasos_enganche,
    "Modificación de masas / MMA":                          _pasos_masas,
    "Modificación de carrocería / Estructura":              _pasos_carroceria,
    "Modificación del sistema de frenos":                   _pasos_frenos_mod,
}


def _desarrollo_calculo(doc, tipo, campos):
    """Escribe el desarrollo paso a paso del cálculo en el documento."""
    fn = _PASOS_MAP.get(tipo)
    if not fn:
        return
    _heading(doc, "Desarrollo del cálculo", level=2, color=COLOR_GRIS)
    fn(doc, campos)


def _seccion_calculo(doc, idx, bloque, reformas_map):
    """Escribe la sección de un bloque de cálculo en el documento."""
    tipo       = bloque.get("tipo", "—")
    reforma_ref = bloque.get("reforma_ref", "— General —")
    campos     = bloque.get("campos", {})

    # Buscar info de la reforma asociada
    reforma_info = reformas_map.get(reforma_ref, {})
    titulo_ref   = reforma_info.get("titulo", "")
    prev_ref     = reforma_info.get("estado_previo", "")
    post_ref     = reforma_info.get("estado_posterior", "")

    # ── Encabezado de sección ─────────────────────────────────────────────
    _heading(doc, f"ANEXO {idx:02d} — {tipo}", level=1, color=COLOR_AZUL)
    _hr(doc)

    # ── Reforma asociada ──────────────────────────────────────────────────
    if reforma_ref and reforma_ref != "— General —":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Reforma asociada: ", bold=True, size=9, color=COLOR_GRIS)
        _run(p, reforma_ref, size=9)

    if titulo_ref:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Título de la reforma: ", bold=True, size=9, color=COLOR_GRIS)
        _run(p, titulo_ref, size=9)

    # Estado previo / posterior (si disponibles)
    if prev_ref or post_ref:
        doc.add_paragraph()  # espacio
        _heading(doc, "Descripción de la reforma", level=2, color=COLOR_GRIS)

        filas_desc = []
        if prev_ref:
            filas_desc.append(("Estado previo",    prev_ref))
        if post_ref:
            filas_desc.append(("Estado posterior", post_ref))
        _tabla_dos_col(doc, filas_desc, widths=(4.5, 12.0))

    doc.add_paragraph()  # espacio

    # ── Tabla de parámetros de entrada y resultados ───────────────────────
    _heading(doc, "Parámetros de cálculo y resultados", level=2, color=COLOR_GRIS)

    # Construir filas: omitir separadores y campos vacíos
    filas_calc = []
    grupos     = []   # para detectar grupos semánticos (separadores)
    grupo_act  = ""

    for clave, valor in campos.items():
        if _es_sep(clave):
            # El valor de un separador es su etiqueta de grupo
            grupo_act = valor if valor else clave.replace("_SEP_", "").replace("_", " ")
            grupos.append(("__SEP__", grupo_act))
        else:
            etiqueta = ETIQUETAS.get(clave, clave)
            filas_calc.append((clave, etiqueta, str(valor), grupo_act))

    # Crear tabla: GRUPO | PARÁMETRO | VALOR
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    hdr = table.rows[0].cells
    hdr[0].width = Cm(3.5)
    hdr[1].width = Cm(8.0)
    hdr[2].width = Cm(5.0)
    for cell, txt in zip(hdr, ["Grupo", "Parámetro", "Valor"]):
        _set_cell_bg(cell, COLOR_FONDO_H)
        p = cell.paragraphs[0]
        _run(p, txt, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Filas de datos
    alt = False
    grupo_prev = ""
    for clave, etiqueta, valor, grupo in filas_calc:
        if not etiqueta or not clave:
            continue
        row = table.add_row()
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(8.0)
        row.cells[2].width = Cm(5.0)

        # Alternar color de fila
        bg = COLOR_FONDO_S if alt else "FFFFFF"
        alt = not alt

        # Si cambia el grupo, usar fondo diferente para la celda de grupo
        grp_txt = grupo if grupo != grupo_prev else ""
        grupo_prev = grupo

        for cell in row.cells:
            _set_cell_bg(cell, bg)

        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p2 = row.cells[2].paragraphs[0]

        _run(p0, grp_txt,  italic=True, size=8,  color=COLOR_GRIS)
        _run(p1, etiqueta, bold=True,   size=9,  color=COLOR_GRIS)

        # Color especial en resultados
        if clave in CAMPOS_RESULTADO:
            ok = "✓" in valor or "OK" in valor.upper()
            col_val = COLOR_VERDE if ok else COLOR_ROJO
            _run(p2, valor, bold=True, size=9, color=col_val)
        elif clave in CAMPOS_COEF:
            try:
                num = float(valor.replace(",", "."))
                col_val = COLOR_VERDE if num <= 1.0 else COLOR_ROJO
            except ValueError:
                col_val = COLOR_GRIS
            _run(p2, valor, bold=True, size=9, color=col_val)
        else:
            _run(p2, valor, size=9)

    doc.add_paragraph()  # espacio tras tabla

    # ── Desarrollo paso a paso del cálculo ──────────────────────────────
    _desarrollo_calculo(doc, tipo, campos)

    doc.add_paragraph()  # espacio tras desarrollo

    # ── Notas y referencia normativa ──────────────────────────────────────
    norma = (campos.get("TORN_NORMA") or campos.get("NORMA_SUSPE")
             or campos.get("MOTOR_NORMA") or campos.get("FLEX_NORMA")
             or campos.get("TORS_NORMA") or campos.get("PAND_NORMA")
             or campos.get("VUELT_NORMA") or campos.get("SOLD_NORMA")
             or campos.get("UA_NORMA") or campos.get("NORMA_FRENOS") or "")
    notas = (campos.get("TORN_NOTAS") or campos.get("NOTAS_SUSPE")
             or campos.get("MOTOR_NOTAS") or "")

    if norma:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Norma / referencia: ", bold=True, size=9, color=COLOR_GRIS)
        _run(p, norma, italic=True, size=9)

    if notas:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Notas: ", bold=True, size=9, color=COLOR_GRIS)
        _run(p, notas, size=9)


# ════════════════════════════════════════════════════════════════════════════
#  Función principal
# ════════════════════════════════════════════════════════════════════════════

def generar_anexo(json_path: Path) -> Path:
    with open(json_path, encoding="utf-8") as f:
        datos = json.load(f)

    calculos = datos.get("CALCULOS", [])
    reformas = datos.get("REFORMAS", [])

    if not calculos:
        raise ValueError("El JSON no contiene bloques de cálculo ('CALCULOS').")

    # Construir mapa reforma_ref -> info (para enriquecer cada sección de cálculo)
    reformas_map = {}
    for r in reformas:
        cod    = r.get("codigo", "")
        titulo = r.get("titulo", "")
        desc   = r.get("descripcion", "")
        if cod:
            # Intentar reconstruir la etiqueta que usa _actualizar_combos_reforma
            if titulo:
                key = f"CR {cod} — {titulo[:50]}"
            elif desc:
                key = f"CR {cod} — {desc[:40]}"
            else:
                key = f"CR {cod}"
            reformas_map[key] = r

    # ── Crear documento ───────────────────────────────────────────────────
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(MARGEN_CM)
        section.bottom_margin = Cm(MARGEN_CM)
        section.left_margin   = Cm(MARGEN_CM + 0.5)
        section.right_margin  = Cm(MARGEN_CM)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Cabecera global
    _cabecera_documento(doc, datos)

    # Una sección por bloque de cálculo
    for i, bloque in enumerate(calculos, start=1):
        if i > 1:
            _add_page_break(doc)
        _seccion_calculo(doc, i, bloque, reformas_map)

    # ── Guardar ───────────────────────────────────────────────────────────
    ref = datos.get("REFERENCIA", "ANEXO").replace("/", "-").replace("\\", "-")
    carpeta_salida = json_path.parent / "proyectos_generados"
    carpeta_salida.mkdir(exist_ok=True)

    nombre = f"ANEXO_{ref}.docx"
    ruta   = carpeta_salida / nombre
    doc.save(str(ruta))
    return ruta


# ════════════════════════════════════════════════════════════════════════════
#  Generación de plantillas de cálculo
# ════════════════════════════════════════════════════════════════════════════

# Campos placeholder por tipo (valores simbólicos para la plantilla)
_PLACEHOLDER_MAP = {
    "Uniones atornilladas — Interior": {
        "TORN_DIAM": "M14", "TORN_CALIDAD": "8.8", "TORN_NUM": "4",
        "TORN_TIPO_CORTE": "Cortante simple", "TORN_A_S": "As", "TORN_A_BRUTA": "A",
        "TORN_FYB": "fyb", "TORN_FUB": "fub", "TORN_FV_RD_1": "Fv,Rd₁",
        "TORN_FV_RD": "Fv,Rd", "TORN_FT_RD": "Ft,Rd",
        "TORN_FV_ED": "Fv,Ed", "TORN_FT_ED": "Ft,Ed",
        "TORN_COEF_CORT": "≤1,0", "TORN_COEF_TRAC": "≤1,0",
        "TORN_INTER": "≤1,0", "TORN_RESULTADO": "✓ / ✗",
    },
    "Flexión simple de viga — Navier": {
        "FLEX_MATERIAL": "S275", "FLEX_FY_MPA": "275", "FLEX_PERFIL": "IPE200",
        "FLEX_W_CM3": "W", "FLEX_M_NM": "M", "FLEX_SIGMA": "σ",
        "FLEX_FD": "fd", "FLEX_COEF": "≤1,0", "FLEX_RESULT": "✓ / ✗",
    },
    "Torsión de perfil — Saint-Venant": {
        "TORS_TIPO": "Cajón cerrado", "TORS_R_MM": "r", "TORS_IP_CM4": "Ip",
        "TORS_FY_MPA": "fy", "TORS_MT_NM": "MT", "TORS_TAU": "τ",
        "TORS_TAU_ADM": "τ_adm", "TORS_COEF": "≤1,0", "TORS_RESULT": "✓ / ✗",
    },
    "Pandeo de barra — Euler": {
        "PAND_MATERIAL": "S275", "PAND_FY_MPA": "275", "PAND_E_MPA": "210000",
        "PAND_L_MM": "L", "PAND_I_CM4": "I", "PAND_A_CM2": "A",
        "PAND_COND": "Biempotrada", "PAND_MU": "μ", "PAND_IRG": "i",
        "PAND_ESBELT": "λ", "PAND_PCR": "Pcr", "PAND_N_N": "N",
        "PAND_COEF": "≤1,0", "PAND_RESULT": "✓ / ✗",
    },
    "Uniones soldadas — cordón en ángulo": {
        "SOLD_MATERIAL": "S275", "SOLD_FU_MPA": "fu", "SOLD_BETA": "βw",
        "SOLD_GARG": "a", "SOLD_LONG": "Lw", "SOLD_F_N": "F",
        "SOLD_TAU": "τ", "SOLD_FW_RD": "fw,Rd",
        "SOLD_COEF": "≤1,0", "SOLD_RESULT": "✓ / ✗",
    },
    "Estabilidad al vuelco lateral": {
        "VUELT_ANCHO": "s", "VUELT_HCG": "hcg", "VUELT_ETA": "η",
        "VUELT_ALIM": "a_lat", "VUELT_VLIM": "V_lim", "VUELT_RESULT": "✓ / ✗",
    },
    "Suspensión neumática — Balonas / Airbag": {
        "MARCA_BALONAS": "MARCA", "PRESION_BALONAS_BAR": "P_máx",
        "NUM_BALONAS": "N", "DIAM_BALONA_MM": "d", "CARGA_TOTAL_KG": "Q",
        "SN_CARGA_BALONA": "q", "SN_AREA_EFECT": "A_ef", "SN_FUERZA_N": "F",
        "SN_P_TRABAJO": "P_t", "SN_MARGEN": "M%", "SN_RESULTADO": "✓ / ✗",
    },
    "Tacos de elevación — Suspensión / Altura": {
        "TACO_MAT": "EPDM", "TACO_MTMA": "MTMA", "TACO_N": "n",
        "TACO_D_MM": "d", "TACO_RCOMP": "R_comp", "TACO_AR": "Ar",
        "TACO_SIGMA": "σ", "TACO_COEF": "≤1,0", "TACO_RESULT": "✓ / ✗",
    },
    "Uniones atornilladas — Exterior (ambiente agresivo)": {
        "TORN_DIAM": "M14", "TORN_CALIDAD": "8.8", "TORN_NUM": "4",
        "TORN_TIPO_CORTE": "Cortante simple", "TORN_A_S": "As", "TORN_A_BRUTA": "A",
        "TORN_FYB": "fyb", "TORN_FUB": "fub", "TORN_FV_RD_1": "Fv,Rd₁",
        "TORN_FV_RD": "Fv,Rd", "TORN_FT_RD": "Ft,Rd",
        "TORN_FV_ED": "Fv,Ed", "TORN_FT_ED": "Ft,Ed",
        "TORN_COEF_CORT": "≤1,0", "TORN_COEF_TRAC": "≤1,0",
        "TORN_INTER": "≤1,0", "TORN_RESULTADO": "✓ / ✗",
        "TORN_EXT_CLASE_CORR": "C3", "TORN_EXT_ACABADO": "Galvanizado",
    },
    "Distribución de cargas por ejes": {
        "EJES_BATALL": "L", "EJES_P_KG": "P", "EJES_D_MM": "d",
        "EJES_Q1_KG": "Q₁", "EJES_Q2_KG": "Q₂",
        "EJES_MMA1": "MMA₁", "EJES_MMA2": "MMA₂",
        "EJES_DQ1": "ΔQ₁", "EJES_DQ2": "ΔQ₂",
        "EJES_NQ1": "Q₁'", "EJES_NQ2": "Q₂'",
        "EJES_RES1": "✓ / ✗", "EJES_RES2": "✓ / ✗",
    },
    "Grúa autocarga — momento de vuelco": {
        "GRUA_CAP_KG": "P", "GRUA_ALC_MM": "L", "GRUA_BCIL_MM": "b_cil",
        "GRUA_MV_NM": "Mv", "GRUA_FCIL_N": "F_cil",
    },
    "Balance de masas — Turismo (M1/M2)": {
        "BM_MMA": "MMA", "BM_MMA_E1": "MMA₁", "BM_MMA_E2": "MMA₂",
        "BM_MT1": "mT₁", "BM_MT2": "mT₂",
        "BM_NP1": "2", "BM_NP2": "3",
        "BM_DT2": "dT₂", "BM_DP1": "dP₁", "BM_DP2": "dP₂",
        "BM_DVT": "dVt", "BM_CAJA": "0",
        "BM_TARA": "Tara", "BM_MQU": "mQu", "BM_DQU": "dQu",
        "BM_R1_SB": "R₁", "BM_R2_SB": "R₂",
        "BM_VER1_SB": "✓ / ✗", "BM_VER2_SB": "✓ / ✗",
    },
    "Unión atornillada — Carga aerodinámica": {
        "UA_CX": "Cx", "UA_V_KMH": "V", "UA_ANCHO": "b", "UA_ALTO": "h",
        "UA_PP_KG": "Pp", "UA_A_M2": "A", "UA_FX_N": "Fx", "UA_FC_N": "Fc",
        "UA_METRICA": "M8", "UA_CALIDAD": "8.8", "UA_CHAPA": "Acero",
        "UA_TMIN": "t", "UA_NT_TRAC": "Nt", "UA_NT_CORT": "Nc",
        "UA_NT_APC": "Na", "UA_NREQ": "N", "UA_RESULT": "✓ / ✗",
    },
    "Unión adhesiva — Carga aerodinámica": {
        "ADH_CX": "Cx", "ADH_V_KMH": "V", "ADH_ANCHO": "b", "ADH_ALTO": "h",
        "ADH_PP_KG": "Pp", "ADH_A_M2": "A", "ADH_FX_N": "Fx", "ADH_FC_N": "Fc",
        "ADH_TIPO": "Sikaflex", "ADH_R_MPA": "R", "ADH_B_MM": "b",
        "ADH_L_MM": "l", "ADH_TAU": "τ", "ADH_COEF": "≤1,0",
        "ADH_RESULT": "✓ / ✗",
    },
    "Protección trasera — barra antichoque": {
        "PT_MTMA": "MTMA", "PT_LONG": "L", "PT_W_CM3": "W", "PT_MAT": "S275",
        "PT_F_HALF": "F", "PT_M_ED": "M_Ed", "PT_SADM": "σ_adm",
        "PT_M_RD": "M_Rd", "PT_COEF": "≤1,0", "PT_RESULT": "✓ / ✗",
    },
    "Frenos — disco (comprobación eficacia)": {
        "FR_MMA": "MMA", "FR_MU": "μ", "FR_ASPECTO": "asp", "FR_SECCION": "sec",
        "FR_LLANTA": "R_ll", "FR_RLLANTA": "R",
        "FR_DEXT_D": "De_d", "FR_DINT_D": "Di_d", "FR_LPAST_D": "Lp_d",
        "FR_DEXT_T": "De_t", "FR_DINT_T": "Di_t", "FR_LPAST_T": "Lp_t",
        "FR_DPIST_D": "Dp_d", "FR_DPIST_T": "Dp_t", "FR_P_MPa": "P",
        "FR_T_DEL": "T_d", "FR_T_TRA": "T_t", "FR_T_TOT": "T_tot",
        "FR_F_FREN": "F_f", "FR_EFIC": "η%", "FR_RATIO50": "r50",
        "FR_RESULT": "✓ / ✗",
    },
    "Suspensión mecánica — Sustitución muelles / amortiguadores": {
        "MARCA_MUELLE_NUEVO": "Marca", "MODELO_MUELLE_NUEVO": "Modelo",
        "MARCA_AMORT_NUEVO": "Marca", "MODELO_AMORT_NUEVO": "Modelo",
        "ALTURA_ANTES_MM": "h₁", "ALTURA_DESPUES_MM": "h₂",
        "SM_VAR_CALC": "Δh", "SM_PCT_VAR": "Δ%",
    },
    "Cambio de motor / Sustitución unidad motriz": {
        "MOTOR_ORIG_TIPO": "Diesel", "MOTOR_ORIG_CILINDRADA_CC": "cc₁",
        "MOTOR_ORIG_POT_KW": "P₁", "MOTOR_ORIG_PAR_NM": "T₁",
        "MOTOR_NUEVO_MARCA": "Marca", "MOTOR_NUEVO_MODELO": "Modelo",
        "MOTOR_NUEVO_CILINDRADA_CC": "cc₂", "MOTOR_NUEVO_POT_KW": "P₂",
        "MOTOR_NUEVO_PAR_NM": "T₂", "MOTOR_TARA_KG": "Tara",
        "RATIO_POT_PESO": "R", "INCREMENTO_POT_PCT": "ΔP%",
        "INCREMENTO_PAR_PCT": "ΔT%", "MOTOR_RESULTADO": "Calculado",
    },
    "Conversión eléctrica o híbrida": {
        "MOTOR_ELEC_MARCA": "Marca", "MOTOR_ELEC_MODELO": "Modelo",
        "MOTOR_ELEC_POT_KW": "P_kW", "BATERIA_MARCA": "Marca",
        "BATERIA_CAPACIDAD_KWH": "C_kWh", "BATERIA_TENSION_V": "V",
        "CONV_TARA_KG": "Tara", "CONV_RATIO_POT": "R", "CONV_ENERGIA_PESO": "E/m",
    },
    "Enganche / Dispositivo de acoplamiento": {
        "ENGANCHE_MARCA": "Marca", "ENGANCHE_TIPO": "Fijo",
        "ENGANCHE_HOMOLOG": "N°", "ENG_MMA_VEHICULO": "MMA_t",
        "MASA_REMOLCABLE_KG": "M_r", "MMA_CONJUNTO_KG": "MMA_c",
        "ENG_MMA_CONJ_CALC": "MMA_calc", "ENG_VERIF_CONJ": "✓ / ✗",
    },
    "Modificación de masas / MMA": {
        "MMA_ANTERIOR_KG": "MMA₁", "MMA_NUEVA_KG": "MMA₂",
        "TARA_KG": "Tara", "CARGA_EJE1_KG": "E₁", "CARGA_EJE2_KG": "E₂",
        "CARGA_EJE3_KG": "E₃", "MMA_MAX_LEGAL_KG": "MMA_leg",
        "MM_SUMA_EJES": "ΣE", "MM_CARGA_UTIL": "CU",
        "MM_VERIF_MMA": "✓ / ✗", "VERIFICACION_EJES": "✓ / ✗",
    },
    "Modificación de carrocería / Estructura": {
        "MATERIAL_CARROCERIA": "Acero",
        "LONG_ANTES_MM": "L₁", "LONG_DESPUES_MM": "L₂",
        "ANCHO_ANTES_MM": "A₁", "ANCHO_DESPUES_MM": "A₂",
        "ALTO_ANTES_MM": "H₁", "ALTO_DESPUES_MM": "H₂",
        "VOLADIZO_DEL_ANTES_MM": "Vd₁", "VOLADIZO_DEL_DESP_MM": "Vd₂",
        "VOLADIZO_TRA_ANTES_MM": "Vt₁", "VOLADIZO_TRA_DESP_MM": "Vt₂",
        "CAR_VAR_LONG": "ΔL", "CAR_VAR_ANCHO": "ΔA", "CAR_VAR_ALTO": "ΔH",
        "CAR_VAR_VOL_DEL": "ΔVd", "CAR_VAR_VOL_TRA": "ΔVt",
    },
    "Modificación del sistema de frenos": {
        "FRENO_SISTEMA_TIPO": "Disco", "FRENO_MARCA_COMP": "Marca",
        "FRENO_MODELO_COMP": "Modelo", "DIST_FRENADA_ANTES": "d₁",
        "DIST_FRENADA_DESP": "d₂", "FRMOD_MEJORA_PCT": "Δ%",
        "FRMOD_RESULTADO": "✓ / ✗",
    },
}


def generar_plantilla_tipo(tipo, carpeta):
    """Genera un documento Word plantilla para un tipo de cálculo.
    Incluye fórmulas OMML con valores simbólicos y secciones de notas editables.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(MARGEN_CM)
        section.bottom_margin = Cm(MARGEN_CM)
        section.left_margin   = Cm(MARGEN_CM + 0.5)
        section.right_margin  = Cm(MARGEN_CM)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Cabecera
    _heading(doc, f"PLANTILLA DE CÁLCULO", level=1, color=COLOR_AZUL)
    _heading(doc, tipo, level=2, color=COLOR_GRIS)
    _hr(doc)

    p_nota = doc.add_paragraph()
    p_nota.paragraph_format.space_before = Pt(6)
    _run(p_nota, "Esta plantilla muestra las fórmulas y pasos de cálculo para este tipo. "
                 "Puede añadir sus comentarios en las secciones marcadas.",
         size=9, italic=True, color=COLOR_GRIS)

    doc.add_paragraph()

    # Generar pasos con valores placeholder
    campos = _PLACEHOLDER_MAP.get(tipo, {})
    fn = _PASOS_MAP.get(tipo)
    if fn and campos:
        _heading(doc, "Desarrollo del cálculo", level=2, color=COLOR_GRIS)
        fn(doc, campos)
    else:
        _descripcion(doc, "Este tipo de cálculo no tiene desarrollo de fórmulas predefinido.")

    # Sección de notas del técnico
    doc.add_paragraph()
    _hr(doc, "FFD54F")
    p_notas = doc.add_paragraph()
    p_notas.paragraph_format.space_before = Pt(8)
    _run(p_notas, "NOTAS DEL TÉCNICO", bold=True, size=11, color=COLOR_AMBAR)

    p_edit = doc.add_paragraph()
    p_edit.paragraph_format.space_before = Pt(4)
    p_edit.paragraph_format.space_after  = Pt(4)
    p_edit.paragraph_format.left_indent  = Cm(0.5)
    _run(p_edit, "[Añada aquí sus observaciones, comentarios o justificaciones adicionales]",
         size=10, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    _hr(doc, "FFD54F")

    # Guardar
    carpeta.mkdir(parents=True, exist_ok=True)
    nombre_safe = tipo.replace("/", "-").replace("\\", "-").replace("—", "-")
    nombre_safe = nombre_safe.replace("  ", " ").strip()
    ruta = carpeta / f"PLANTILLA_{nombre_safe}.docx"
    doc.save(str(ruta))
    return ruta


def generar_todas_plantillas(carpeta):
    """Genera una plantilla Word por cada tipo de cálculo con fórmulas."""
    resultados = []
    for tipo in _PASOS_MAP:
        try:
            ruta = generar_plantilla_tipo(tipo, carpeta)
            resultados.append((tipo, str(ruta), True))
        except Exception as e:
            resultados.append((tipo, str(e), False))
    return resultados


# ════════════════════════════════════════════════════════════════════════════
#  Entry-point
# ════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python generar_anexo.py <ruta_json>", file=sys.stderr)
        sys.exit(1)

    json_path = Path(sys.argv[1])
    if not json_path.exists():
        print(f"Archivo no encontrado: {json_path}", file=sys.stderr)
        sys.exit(1)

    try:
        ruta = generar_anexo(json_path)
        print(str(ruta))          # formulario.py captura esta línea
        sys.exit(0)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
